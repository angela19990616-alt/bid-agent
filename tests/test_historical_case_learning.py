import json
from io import BytesIO
from uuid import uuid4

from docx import Document

from app.knowledge.permissions import KnowledgeAccessContext
from app.memory.historical_case_learning import (
    HistoricalCaseLearningService,
    HistoricalCasePair,
    HistoricalCasePatternExtractor,
)


def _winning_bid() -> bytes:
    document = Document()
    document.add_heading("第一章 项目理解", level=1)
    document.add_paragraph("甲公司为某市客户建设了金额500万元的智慧平台。")
    document.add_heading("1.1 总体思路", level=2)
    document.add_paragraph("采用分层方法展开论证，并给出闭环措施。")
    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "阶段"
    table.cell(0, 1).text = "时间"
    table.cell(0, 2).text = "成果"
    table.cell(1, 0).text = "历史阶段A"
    table.cell(1, 1).text = "30天"
    table.cell(1, 2).text = "旧项目成果"
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _five_level_winning_bid() -> bytes:
    document = Document()
    for level, title in enumerate([
        "第一章 项目理解",
        "1.1 总体思路",
        "1.1.1 实施路径",
        "1.1.1.1 质量控制",
        "1.1.1.1.1 检查记录",
    ], start=1):
        document.add_heading(title, level=level)
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_case_learning_keeps_structure_but_removes_historical_facts():
    patterns = HistoricalCasePatternExtractor().extract(_winning_bid())
    snapshot = json.dumps(patterns, ensure_ascii=False)

    assert patterns
    assert patterns[0]["prohibited_fact_copy"] is True
    assert patterns[0]["source_facts_removed"] is True
    assert "项目理解" in snapshot
    assert "阶段" in snapshot and "时间" in snapshot and "成果" in snapshot
    assert "甲公司" not in snapshot
    assert "某市客户" not in snapshot
    assert "500万元" not in snapshot
    assert "历史阶段A" not in snapshot


def test_case_learning_preserves_five_level_structure_depth():
    patterns = HistoricalCasePatternExtractor().extract(
        _five_level_winning_bid()
    )

    assert max(item["visual_pattern"]["heading_level"] for item in patterns) == 5


def test_case_learning_rejects_numbered_body_sentences_as_headings():
    document = Document()
    document.add_paragraph("一、整体服务方案")
    document.add_paragraph("1.我方承诺响应采购文件的全部要求。")
    document.add_paragraph("2.本表可根据实际情况进行拓展。")
    document.add_paragraph("二、质量保证措施")
    stream = BytesIO()
    document.save(stream)

    patterns = HistoricalCasePatternExtractor().extract(stream.getvalue())

    assert [item["chapter_structure"][0] for item in patterns] == [
        "总体思路",
        "质量保障",
    ]


def test_case_learning_maps_person_specific_heading_without_personal_fact():
    document = Document()
    document.add_paragraph("1.项目负责人-某员工")
    stream = BytesIO()
    document.save(stream)

    patterns = HistoricalCasePatternExtractor().extract(stream.getvalue())
    snapshot = json.dumps(patterns, ensure_ascii=False)

    assert any(
        item["chapter_structure"][0] == "人员材料" for item in patterns
    )
    assert "某员工" not in snapshot


def test_case_learning_drops_unknown_numbered_source_specific_titles():
    document = Document()
    document.add_paragraph("一、类似项目的成功案例")
    document.add_paragraph("1.某地历史项目专有名称")
    stream = BytesIO()
    document.save(stream)

    patterns = HistoricalCasePatternExtractor().extract(stream.getvalue())
    snapshot = json.dumps(patterns, ensure_ascii=False)

    assert {
        item["chapter_structure"][0] for item in patterns
    } == {"方案正文", "业绩证明"}
    assert "某地历史项目专有名称" not in snapshot


def test_case_pair_learning_preserves_private_organization_boundary():
    calls = []

    class FakeMemoryEngine:
        def add_pattern(self, **kwargs):
            calls.append(kwargs)
            return uuid4()

    context = KnowledgeAccessContext("enterprise-a")
    learned = HistoricalCaseLearningService(
        memory_engine=FakeMemoryEngine()
    ).learn_pairs(
        access_context=context,
        pairs=[
            HistoricalCasePair(
                tender_filename="tender.docx",
                tender_text="智慧文旅项目采购需求",
                proposal_filename="winning.docx",
                proposal_content=_winning_bid(),
                project_type="咨询服务",
                industry="文旅",
            )
        ],
    )

    assert learned
    assert calls
    assert all(item["access_context"] is context for item in calls)
    assert all(
        item["pattern"]["reference_scope"] == "organization_private"
        for item in calls
    )
    snapshot = json.dumps(
        [item["pattern"] for item in calls], ensure_ascii=False
    )
    assert "甲公司" not in snapshot
    assert "500万元" not in snapshot


def test_multiple_pairs_use_one_validated_memory_transaction():
    calls = []

    class FakeMemoryEngine:
        def add_patterns(self, **kwargs):
            calls.append(kwargs)
            return [uuid4() for _ in kwargs["items"]]

    pairs = [
        HistoricalCasePair(
            tender_filename=f"tender-{index}.docx",
            tender_text=f"第{index}组采购需求",
            proposal_filename=f"winning-{index}.docx",
            proposal_content=_winning_bid(),
            project_type="咨询服务",
            industry="文旅",
        )
        for index in range(1, 6)
    ]
    learned = HistoricalCaseLearningService(
        memory_engine=FakeMemoryEngine()
    ).learn_pairs(
        access_context=KnowledgeAccessContext("enterprise-a"),
        pairs=pairs,
    )

    assert len(calls) == 1
    assert calls[0]["replace_source_pairs"] is True
    assert len(learned) == len(calls[0]["items"])
    assert all(
        item["pattern"]["source_facts_removed"] is True
        and item["pattern"]["prohibited_fact_copy"] is True
        for item in calls[0]["items"]
    )
