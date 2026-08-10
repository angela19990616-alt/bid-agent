from io import BytesIO
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from app.services.response_template_service import ResponseTemplateService


def _template_bytes(*, embedded: bool = True) -> bytes:
    document = Document()
    if embedded:
        document.add_heading("采购需求", level=1)
        document.add_paragraph("此处是采购文件正文，不应进入响应文件。")
    document.add_heading("附件：响应文件格式", level=1)
    document.add_paragraph("投标人必须严格按照本格式填写，不得修改表格。")
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "项目名称"
    table.cell(0, 1).text = "{{project_name}}"
    table.cell(1, 0).text = "供应商名称"
    table.cell(1, 1).text = ""
    document.add_heading("技术方案", level=1)
    document.add_paragraph("请在此处填写。")
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_detects_embedded_docx_template_and_strict_rules():
    descriptor = ResponseTemplateService().detect(
        "招标文件.docx", _template_bytes()
    )

    assert descriptor.detected is True
    assert descriptor.template_kind == "embedded_response_section"
    assert descriptor.fidelity == "exact_template"
    assert descriptor.start_block is not None
    assert "project_name" in descriptor.placeholders
    assert "严格按照" in descriptor.strict_reasons
    assert descriptor.semantic_audit is not None
    assert descriptor.semantic_audit["status"] == "passed"
    assert descriptor.semantic_audit["issue_count"] == 0


def test_docx_negated_template_statement_does_not_trigger_strict_mode():
    document = Document()
    document.add_heading("采购需求", level=1)
    document.add_paragraph("本项目未提供统一的投标文件格式，投标人自行编制技术方案。")
    stream = BytesIO()
    document.save(stream)

    descriptor = ResponseTemplateService().detect(
        "招标文件.docx", stream.getvalue()
    )

    assert descriptor.detected is False
    assert descriptor.fidelity == "planned"


def test_pdf_negated_template_statement_does_not_claim_reference():
    class FakePage:
        @staticmethod
        def extract_text():
            return "本项目没有提供投标文件格式，投标人自行编制。"

    class FakeReader:
        pages = [FakePage()]

    with patch(
        "app.services.response_template_service.PdfReader",
        return_value=FakeReader(),
    ):
        descriptor = ResponseTemplateService().detect("招标文件.pdf", b"pdf")

    assert descriptor.detected is False
    assert descriptor.fidelity == "manual_exact_fill_required"


def test_extracts_response_template_outline_up_to_five_levels():
    document = Document()
    document.add_heading("采购需求", level=1)
    document.add_heading("附件：投标文件格式", level=1)
    document.add_heading("一、投标函", level=1)
    document.add_heading("（一）项目基本信息", level=2)
    document.add_heading("1.1 报价信息", level=3)
    document.add_heading("1.1.1 报价明细", level=4)
    document.add_heading("1.1.1.1 其他说明", level=5)
    document.add_heading(
        "一、本段是很长的说明文字，不应当作为可编辑目录标题展示给用户，而应当保留在原模板正文中。",
        level=1,
    )
    document.add_heading("第八章 评审办法", level=1)
    document.add_heading("一、评审原则", level=1)
    stream = BytesIO()
    document.save(stream)

    descriptor = ResponseTemplateService().detect(
        "招标文件.docx", stream.getvalue()
    )

    assert [item["title"] for item in descriptor.outline] == [
        "附件：投标文件格式",
        "一、投标函",
        "（一）项目基本信息",
        "1.1 报价信息",
        "1.1.1 报价明细",
        "1.1.1.1 其他说明",
    ]
    assert [item["level"] for item in descriptor.outline] == [1, 1, 2, 3, 4, 5]


def test_repeated_response_format_title_prefers_later_real_chapter():
    document = Document()
    document.add_paragraph("第六章 响应文件格式")
    document.add_heading("三、响应文件的编制", level=1)
    document.add_paragraph(
        "说明：本部分格式用于询问函，不属于响应文件格式的组成部分。"
    )
    document.add_heading("第五章 合同条款", level=1)
    document.add_paragraph("第六章  响应文件格式")
    document.add_heading("一、响应函", level=1)
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "供应商名称"
    table.cell(0, 1).text = ""
    stream = BytesIO()
    document.save(stream)

    descriptor = ResponseTemplateService().detect(
        "采购文件.docx", stream.getvalue()
    )

    assert descriptor.marker_text == "第六章  响应文件格式"
    assert [item["title"] for item in descriptor.outline] == [
        "第六章 响应文件格式",
        "一、响应函",
    ]
    assert any(
        item["field_key"] == "bidder_name"
        for item in descriptor.fields
    )


def test_fills_verified_values_and_keeps_original_table(tmp_path):
    service = ResponseTemplateService()
    content = _template_bytes()
    descriptor = service.detect("招标文件.docx", content)
    output = tmp_path / "filled.docx"

    report = service.fill_docx(
        template_content=content,
        output_path=output,
        descriptor=descriptor,
        field_values={
            "project_name": "测试采购项目",
            "bidder_name": "已核验供应商",
        },
        sections=[{"title": "实施计划", "content": "分三个阶段实施。"}],
    )

    result = Document(output)
    full_text = "\n".join(
        [item.text for item in result.paragraphs]
        + [cell.text for table in result.tables for row in table.rows for cell in row.cells]
    )
    assert "采购文件正文" not in full_text
    assert "测试采购项目" in full_text
    assert "已核验供应商" in full_text
    assert "分三个阶段实施" in full_text
    assert len(result.tables) == 1
    assert report.unresolved_sections == ()


def test_missing_template_value_is_not_invented(tmp_path):
    service = ResponseTemplateService()
    content = _template_bytes(embedded=False)
    descriptor = service.detect("响应文件格式模板.docx", content)
    output = tmp_path / "filled.docx"

    report = service.fill_docx(
        template_content=content,
        output_path=output,
        descriptor=descriptor,
        field_values={},
        sections=[],
    )

    assert "project_name" in report.unresolved_fields
    assert "{{project_name}}" in "\n".join(
        cell.text
        for table in Document(output).tables
        for row in table.rows
        for cell in row.cells
    )


def test_repeated_phone_slots_keep_role_specific_values(tmp_path):
    document = Document()
    document.add_heading("附件：响应文件格式", level=1)
    table = document.add_table(rows=3, cols=3)
    for row_index, (role, value) in enumerate((
        ("联系人", "13800138000"),
        ("法定代表人", "010-12345678"),
        ("技术负责人", "13900139000"),
    )):
        table.cell(row_index, 0).text = role
        table.cell(row_index, 1).text = "联系电话"
        table.cell(row_index, 2).text = ""
    stream = BytesIO()
    document.save(stream)
    content = stream.getvalue()
    service = ResponseTemplateService()
    descriptor = service.detect("响应文件格式.docx", content)
    fields = {
        item["expected_role"]: item
        for item in descriptor.fields
        if item.get("canonical_key") == "contact_phone"
    }

    assert set(fields) == {
        "CONTACT_PERSON", "LEGAL_REPRESENTATIVE", "TECHNICAL_LEAD"
    }
    assert len({item["field_key"] for item in fields.values()}) == 3
    values = {
        fields[role]["field_key"]: value
        for role, value in (
            ("CONTACT_PERSON", "13800138000"),
            ("LEGAL_REPRESENTATIVE", "010-12345678"),
            ("TECHNICAL_LEAD", "13900139000"),
        )
    }
    output = tmp_path / "role-phones.docx"
    service.fill_docx(
        template_content=content,
        output_path=output,
        descriptor=descriptor,
        field_values=values,
        sections=[],
    )

    result = Document(output)
    assert [row.cells[2].text for row in result.tables[0].rows] == [
        "13800138000", "010-12345678", "13900139000"
    ]


def test_pdf_template_is_not_falsely_claimed_as_auto_fillable():
    descriptor = ResponseTemplateService().detect(
        "附件-投标文件格式.pdf", b"%PDF-1.7"
    )

    assert descriptor.detected is True
    assert descriptor.fidelity == "manual_exact_fill_required"


def test_extracts_project_number_from_procurement_source_only():
    document = Document()
    document.add_paragraph("采购项目名称：测试采购项目。")
    document.add_paragraph("项目编号：SCXHR20250320。")
    document.add_paragraph("供应商名称：此处不得从历史响应文件推断")
    stream = BytesIO()
    document.save(stream)

    values = ResponseTemplateService().extract_source_fields(
        "招标文件.docx", stream.getvalue()
    )

    assert values == {
        "project_name": "测试采购项目",
        "project_number": "SCXHR20250320",
    }

    detailed_values, evidence = (
        ResponseTemplateService().extract_source_fields_with_evidence(
            "招标文件.docx", stream.getvalue()
        )
    )
    assert detailed_values == values
    assert evidence["project_number"]["title"] == "招标文件.docx"
    assert evidence["project_name"]["location"] == "正文第 1 段"
    assert evidence["project_number"]["location"] == "正文第 2 段"
    assert "SCXHR20250320" in evidence["project_number"]["excerpt"]


def test_empty_paragraph_is_never_selected_as_section_heading(tmp_path):
    document = Document()
    document.add_paragraph("")
    document.add_heading("附件：响应文件格式", level=1)
    document.add_paragraph("")
    document.add_heading("技术方案", level=1)
    stream = BytesIO()
    document.save(stream)
    content = stream.getvalue()
    service = ResponseTemplateService()
    descriptor = service.detect("响应文件格式.docx", content)

    report = service.fill_docx(
        template_content=content,
        output_path=tmp_path / "filled.docx",
        descriptor=descriptor,
        field_values={},
        sections=[{"title": "技术方案", "content": "应写在本标题后。"}],
    )

    paragraphs = [item.text for item in Document(tmp_path / "filled.docx").paragraphs]
    heading_index = paragraphs.index("技术方案")
    content_index = paragraphs.index("应写在本标题后。")
    assert content_index == heading_index + 1
    assert report.unresolved_sections == ()


def test_technical_solution_does_not_match_short_table_label():
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "技术"
    expected = document.add_paragraph("5.格式13 项目实施方案")

    target = ResponseTemplateService._find_heading(document, "技术方案")

    assert target is not None
    assert target._p is expected._p


def test_form_control_prefix_is_hidden_from_user_facing_outline():
    assert ResponseTemplateService._display_heading(
        "5.格式13 项目实施方案"
    ) == "项目实施方案"


def test_inserted_content_does_not_inherit_heading_page_break():
    document = Document()
    anchor = document.add_paragraph("项目实施方案")
    anchor.paragraph_format.page_break_before = True

    ResponseTemplateService._insert_content_after(anchor, "1. 工作范围\n正文内容")

    stream = BytesIO()
    document.save(stream)
    result = Document(BytesIO(stream.getvalue()))
    heading, body = result.paragraphs[1:3]
    assert heading.paragraph_format.page_break_before is not True
    assert body.paragraph_format.page_break_before is not True
    assert heading.runs[0].bold is True
    assert body.runs[0].font.color.rgb == RGBColor(0, 0, 0)


def test_detects_template_fonts_and_uses_them_for_inserted_content(tmp_path):
    document = Document()
    document.add_heading("附件：响应文件格式", level=1)
    heading = document.add_heading("技术方案", level=1)
    heading_run = heading.runs[0]
    heading_run.font.name = "黑体"
    heading_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    body = document.add_paragraph("正文样式参考")
    body_run = body.runs[0]
    body_run.font.name = "仿宋_GB2312"
    body_run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    body_run.font.size = Pt(16)
    stream = BytesIO()
    document.save(stream)
    content = stream.getvalue()
    service = ResponseTemplateService()
    descriptor = service.detect("投标文件格式.docx", content)

    assert "仿宋_GB2312" in descriptor.font_profile["detected_fonts"]
    assert "黑体" in descriptor.font_profile["detected_fonts"]

    output = tmp_path / "font-preserved.docx"
    service.fill_docx(
        template_content=content,
        output_path=output,
        descriptor=descriptor,
        field_values={},
        sections=[{"title": "技术方案", "content": "一、实施思路\n这是新生成的正文。"}],
    )
    result = Document(output)
    generated_body = next(
        item for item in result.paragraphs if item.text == "这是新生成的正文。"
    )
    fonts = generated_body.runs[0]._element.rPr.rFonts
    assert fonts.get(qn("w:eastAsia")) == "仿宋_GB2312"
    assert generated_body.runs[0].font.size.pt == 16


def test_compound_labels_use_the_value_before_the_blank_and_skip_signatures():
    document = Document()
    document.add_heading("附件：投标文件格式", level=1)
    table = document.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "项目编号/包号：_____________________ 项目名称"
    table.cell(1, 0).text = "企业名称（盖章）"
    table.cell(2, 0).text = "法定代表人（单位负责人）（签字或盖章）"
    table.cell(3, 0).text = "附：法定代表人身份证明文件复印件"
    stream = BytesIO()
    document.save(stream)

    descriptor = ResponseTemplateService().detect(
        "投标文件格式.docx", stream.getvalue()
    )
    by_label = {item["label"]: item["field_key"] for item in descriptor.fields}

    assert by_label["项目编号/包号"] == "project_number"
    assert all("_" not in label and "＿" not in label for label in by_label)
    assert by_label["企业名称（盖章）"] == "bidder_name"
    assert "法定代表人（单位负责人）（签字或盖章）" not in by_label
    assert "附：法定代表人身份证明文件复印件" not in by_label


def test_multiple_labeled_blanks_are_split_into_independent_business_slots(
    tmp_path,
):
    document = Document()
    document.add_heading("附件：投标文件格式", level=1)
    document.add_paragraph("格式9供应商基本情况表")
    document.add_paragraph(
        "项目编号：________________ 项目名称：________________"
    )
    document.add_paragraph(
        "姓名：____ 性别：____ 年龄：____ 职务：____"
    )
    stream = BytesIO()
    document.save(stream)
    content = stream.getvalue()

    service = ResponseTemplateService()
    descriptor = service.detect("投标文件格式.docx", content)
    fields = descriptor.fields

    assert [item["label"] for item in fields[:2]] == [
        "项目编号", "项目名称",
    ]
    assert [item["canonical_key"] for item in fields[:2]] == [
        "project_number", "project_name",
    ]
    assert all("_" not in item["label"] for item in fields)
    assert all(
        item["document_section"].endswith("供应商基本情况表")
        for item in fields
    )

    output = tmp_path / "split-fields.docx"
    service.fill_docx(
        template_content=content,
        output_path=output,
        descriptor=descriptor,
        field_values={
            fields[0]["field_key"]: "P-001",
            fields[1]["field_key"]: "测试项目",
        },
        sections=[],
    )
    rendered = "\n".join(item.text for item in Document(output).paragraphs)
    assert "项目编号：P-001" in rendered
    assert "项目名称：测试项目" in rendered


def test_instruction_paragraph_does_not_replace_the_current_form_directory():
    document = Document()
    document.add_heading("附件：投标文件格式", level=1)
    document.add_paragraph("格式10技术要求应答表")
    document.add_paragraph("2.上述各项可另页描述。")
    document.add_paragraph("项目名称：________________")
    stream = BytesIO()
    document.save(stream)

    descriptor = ResponseTemplateService().detect(
        "投标文件格式.docx", stream.getvalue()
    )

    project_field = next(
        item for item in descriptor.fields
        if item["canonical_key"] == "project_name"
    )
    assert project_field["document_section"].endswith("技术要求应答表")


def test_legal_declaration_ending_with_colon_is_not_a_fill_slot():
    document = Document()
    document.add_heading("第六章 响应文件格式", level=1)
    document.add_paragraph(
        "具备《中华人民共和国政府采购法》第二十二条规定的条件："
    )
    document.add_paragraph("签名代表在此声明并同意：")
    document.add_paragraph("供应商名称：________")
    stream = BytesIO()
    document.save(stream)

    descriptor = ResponseTemplateService().detect(
        "采购文件.docx", stream.getvalue()
    )

    assert [item["canonical_key"] for item in descriptor.fields] == [
        "bidder_name",
    ]


def test_strict_fill_keeps_header_footer_and_sets_delivery_metadata(tmp_path):
    document = Document()
    document.sections[0].header.paragraphs[0].text = "原模板页眉"
    document.sections[0].footer.paragraphs[0].text = "原模板页脚"
    document.add_heading("附件：投标文件格式", level=1)
    document.add_paragraph("项目名称：{{project_name}}")
    stream = BytesIO()
    document.save(stream)
    content = stream.getvalue()
    descriptor = ResponseTemplateService().detect("投标文件格式.docx", content)
    assert all(
        "页眉" not in item["surrounding_text"]
        and "页脚" not in item["surrounding_text"]
        for item in descriptor.fields
    )
    output = tmp_path / "metadata.docx"

    ResponseTemplateService().fill_docx(
        template_content=content,
        output_path=output,
        descriptor=descriptor,
        field_values={"project_name": "测试项目"},
        sections=[],
        document_title="《AI投标文件+测试项目》",
    )

    result = Document(output)
    assert result.core_properties.title == "《AI投标文件+测试项目》"
    assert result.sections[0].header.paragraphs[0].text == "原模板页眉"
    assert result.sections[0].footer.paragraphs[0].text == "原模板页脚"


def test_table_fill_preserves_existing_target_font():
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "供应商名称"
    target = table.cell(0, 1)
    run = target.paragraphs[0].add_run("{{bidder_name}}")
    run.font.name = "仿宋_GB2312"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")

    ResponseTemplateService._set_cell_text_preserving_style(
        target, "已核验企业"
    )

    assert target.text == "已核验企业"
    assert target.paragraphs[0].runs[0]._element.rPr.rFonts.get(
        qn("w:eastAsia")
    ) == "仿宋_GB2312"


def test_fills_visible_blank_label_paragraphs():
    document = Document()
    document.add_paragraph("项目编号： __________")
    document.add_paragraph("项目名称：")
    document.add_paragraph("供应商（加盖公章）：      ")
    filled: set[str] = set()

    ResponseTemplateService._fill_label_paragraphs(
        document,
        {
            "project_number": "SCXHR20250320",
            "project_name": "自贡智慧文旅项目",
            "bidder_name": "【待人工确认供应商名称】",
        },
        filled,
    )

    assert document.paragraphs[0].text == "项目编号： SCXHR20250320"
    assert document.paragraphs[1].text == "项目名称： 自贡智慧文旅项目"
    assert "【待人工确认供应商名称】" in document.paragraphs[2].text
    assert filled == {"project_number", "project_name", "bidder_name"}


def test_detects_and_fills_generic_blank_table_field(tmp_path):
    document = Document()
    document.add_heading("第七章 响应文件格式", level=1)
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "投标报价金额"
    table.cell(0, 1).text = ""
    stream = BytesIO()
    document.save(stream)
    content = stream.getvalue()
    service = ResponseTemplateService()
    descriptor = service.detect("招标文件.docx", content)
    field = next(
        item for item in descriptor.fields
        if item["label"] == "投标报价金额"
    )

    assert field["expected_source"] == "pricing_database"
    assert field["field_key"] == "unmapped_field"
    assert field["display_name"] == "尚未识别的业务槽位"

    output = tmp_path / "generic-fill.docx"
    report = service.fill_docx(
        template_content=content,
        output_path=output,
        descriptor=descriptor,
        field_values={field["field_key"]: "100000.00元"},
        sections=[],
    )

    result = Document(output)
    assert result.tables[0].cell(0, 1).text == "100000.00元"
    assert field["field_key"] in report.filled_fields


def test_detects_empty_table_cells_from_horizontal_and_vertical_labels(tmp_path):
    document = Document()
    document.add_heading("第七章 响应文件格式", level=1)
    horizontal = document.add_table(rows=1, cols=4)
    horizontal.cell(0, 0).text = "姓名"
    horizontal.cell(0, 1).text = ""
    horizontal.cell(0, 2).text = "联系电话"
    horizontal.cell(0, 3).text = ""
    vertical = document.add_table(rows=3, cols=2)
    vertical.cell(0, 0).text = "项目名称"
    vertical.cell(0, 1).text = "项目编号"
    vertical.cell(1, 0).text = ""
    vertical.cell(1, 1).text = ""
    vertical.cell(2, 0).text = ""
    vertical.cell(2, 1).text = ""
    stream = BytesIO()
    document.save(stream)

    descriptor = ResponseTemplateService().detect("表格模板.docx", stream.getvalue())
    fields = {(item["label"], item["source_location"]) for item in descriptor.fields}

    assert any(label == "姓名" and "第1行/第2列" in location for label, location in fields)
    assert any(label == "联系电话" and "第1行/第4列" in location for label, location in fields)
    assert any(label == "项目名称" and "第2行/第1列" in location for label, location in fields)
    assert any(label == "项目编号" and "第2行/第2列" in location for label, location in fields)
    assert any(label == "项目名称" and "第3行/第1列" in location for label, location in fields)
    assert any(label == "项目编号" and "第3行/第2列" in location for label, location in fields)


def test_detects_colon_and_inline_cell_blanks_and_fills_each_slot(tmp_path):
    document = Document()
    document.add_heading("第七章 响应文件格式", level=1)
    document.add_paragraph("投标人名称：")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "法定代表人姓名：____  联系电话：____"
    stream = BytesIO()
    document.save(stream)
    content = stream.getvalue()
    service = ResponseTemplateService()
    descriptor = service.detect("授权书.docx", content)

    bidder = next(item for item in descriptor.fields if item["label"] == "投标人名称")
    legal_name = next(
        item for item in descriptor.fields if item["label"] == "法定代表人姓名"
    )
    phone = next(item for item in descriptor.fields if item["label"] == "联系电话")
    assert bidder["target_mode"] == "inline_paragraph"
    assert legal_name["target_mode"] == "inline_cell"
    assert legal_name["expected_role"] == "LEGAL_REPRESENTATIVE"

    output = tmp_path / "filled.docx"
    service.fill_docx(
        template_content=content,
        output_path=output,
        descriptor=descriptor,
        field_values={
            bidder["field_key"]: "北京示例咨询有限公司",
            legal_name["field_key"]: "张三",
            phone["field_key"]: "13800000000",
        },
        sections=[],
    )
    result = Document(output)
    assert result.paragraphs[1].text == "投标人名称：北京示例咨询有限公司"
    assert result.tables[0].cell(0, 0).text == (
        "法定代表人姓名：张三  联系电话：13800000000"
    )


def test_standalone_parenthesized_document_title_is_not_a_fill_slot():
    document = Document()
    document.add_heading("第七章 响应文件格式", level=1)
    document.add_paragraph("响  应  文  件")
    document.add_paragraph("（资格性响应文件）")
    document.add_paragraph("项目名称：")
    stream = BytesIO()
    document.save(stream)

    descriptor = ResponseTemplateService().detect(
        "资格性响应文件.docx", stream.getvalue()
    )

    assert [item["label"] for item in descriptor.fields] == ["项目名称"]


def test_personnel_grid_uses_column_attribute_not_left_row_category():
    document = Document()
    document.add_heading("第七章 响应文件格式", level=1)
    table = document.add_table(rows=3, cols=4)
    for column, label in enumerate(("类别", "职务", "姓名", "技术职称")):
        table.cell(0, column).text = label
    table.cell(1, 0).text = "管理人员"
    table.cell(2, 0).text = "技术人员"
    stream = BytesIO()
    document.save(stream)

    descriptor = ResponseTemplateService().detect(
        "项目人员情况表.docx", stream.getvalue()
    )
    row_two = {
        item["column"]: item
        for item in descriptor.fields
        if item["row"] == 2
    }

    assert row_two[2]["label"] == "职务"
    assert row_two[2]["semantic_field"] == "person.title"
    assert row_two[3]["label"] == "姓名"
    assert row_two[3]["semantic_field"] == "person.name"
    assert row_two[4]["semantic_field"] == "person.professional_title"
    assert all(
        item["semantic_field"] != "project.team.members"
        for item in descriptor.fields
    )


def test_remark_slot_does_not_inherit_a_person_header_from_prior_rows():
    document = Document()
    document.add_heading("第七章 响应文件格式", level=1)
    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = "联系电话"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = ""
    table.cell(2, 0).text = "备注"
    table.cell(2, 1).text = ""
    stream = BytesIO()
    document.save(stream)

    descriptor = ResponseTemplateService().detect(
        "人员信息表.docx", stream.getvalue()
    )
    remark = next(
        item for item in descriptor.fields
        if item["source_location"].endswith("第3行/第2列")
    )

    assert remark["label"] == "备注"
    assert remark["semantic_field"] == "bid_response.content"
    assert remark["required"] is False


def test_signature_marks_are_actions_not_text_fields():
    document = Document()
    document.add_heading("第七章 响应文件格式", level=1)
    document.add_paragraph("投标人（盖章）：____")
    document.add_paragraph("法定代表人（签字）：____")
    stream = BytesIO()
    document.save(stream)

    descriptor = ResponseTemplateService().detect("签章模板.docx", stream.getvalue())

    assert descriptor.fields == ()
    action_names = {item["display_name"] for item in descriptor.actions}
    assert "加盖投标人公章" in action_names
    assert "法定代表人签字" in action_names


def test_signature_choice_keeps_both_authorized_roles_in_the_action():
    document = Document()
    document.add_heading("第七章 响应文件格式", level=1)
    document.add_paragraph("法定代表人或其委托代理人签字：____")
    stream = BytesIO()
    document.save(stream)

    descriptor = ResponseTemplateService().detect(
        "签章选择.docx", stream.getvalue()
    )

    assert descriptor.fields == ()
    assert [item["display_name"] for item in descriptor.actions] == [
        "法定代表人或授权代表签字"
    ]


def test_real_world_signature_aliases_are_actions_not_people_or_text_fields():
    document = Document()
    document.add_heading("第七章 响应文件格式", level=1)
    document.add_paragraph("法定代表人或授权委托人（签字）：____")
    document.add_paragraph("授权委托代理人签字或盖章：____")
    document.add_paragraph("法定代表人/负责人或被授权委托人（签字或盖章）：____")
    document.add_paragraph("法定代表人/负责人或其授权代表（签字或盖章）：____")
    stream = BytesIO()
    document.save(stream)

    descriptor = ResponseTemplateService().detect(
        "真实签章写法.docx", stream.getvalue()
    )

    assert descriptor.fields == ()
    assert {item["display_name"] for item in descriptor.actions} == {
        "法定代表人或授权代表签字",
        "授权代表签字",
        "加盖投标人公章",
    }


def test_value_slot_with_seal_is_one_bidder_variable_plus_document_action():
    document = Document()
    document.add_heading("第七章 响应文件格式", level=1)
    document.add_paragraph("投标人名称（盖章）：____")
    document.add_paragraph("供应商名称：____")
    stream = BytesIO()
    document.save(stream)

    descriptor = ResponseTemplateService().detect(
        "投标人名称模板.docx", stream.getvalue()
    )

    assert [item["canonical_key"] for item in descriptor.fields] == [
        "bidder_name", "bidder_name",
    ]
    assert {item["display_name"] for item in descriptor.actions} == {
        "加盖投标人公章",
    }


def test_colon_introductions_and_section_headings_are_not_fill_slots():
    document = Document()
    document.add_heading("第七章 响应文件格式", level=1)
    document.add_paragraph("现郑重承诺如下：")
    document.add_paragraph("一、具备本项目规定的条件：")
    document.add_paragraph("投标人名称：")
    stream = BytesIO()
    document.save(stream)

    descriptor = ResponseTemplateService().detect("承诺函.docx", stream.getvalue())
    labels = {item["label"] for item in descriptor.fields}

    assert "投标人名称" in labels
    assert "现郑重承诺如下" not in labels
    assert "一、具备本项目规定的条件" not in labels


def test_new_form_directory_does_not_inherit_previous_form_subheading():
    document = Document()
    document.add_heading("第七章 响应文件格式", level=1)
    document.add_paragraph("格式3 承诺函")
    document.add_paragraph("一、具备本项目规定的条件：")
    document.add_paragraph("投标人名称：____")
    document.add_paragraph("格式9 供应商基本情况表")
    document.add_paragraph("注册地址：____")
    stream = BytesIO()
    document.save(stream)

    descriptor = ResponseTemplateService().detect("目录模板.docx", stream.getvalue())
    bidder = next(item for item in descriptor.fields if item["label"] == "投标人名称")
    address = next(item for item in descriptor.fields if item["label"] == "注册地址")

    assert bidder["document_section"].endswith(
        "格式3 承诺函 / 一、具备本项目规定的条件："
    )
    assert address["document_section"] == (
        "第七章 响应文件格式 / 格式9 供应商基本情况表"
    )


def test_detects_and_fills_generic_blank_paragraph_field(tmp_path):
    document = Document()
    document.add_heading("第七章 响应文件格式", level=1)
    document.add_paragraph("开户行： __________")
    document.add_heading("第八章 合同条款", level=1)
    document.add_paragraph("合同编号： __________")
    stream = BytesIO()
    document.save(stream)
    content = stream.getvalue()
    service = ResponseTemplateService()
    descriptor = service.detect("招标文件.docx", content)
    field = next(item for item in descriptor.fields if item["label"] == "开户行")

    assert all(item["label"] != "合同编号" for item in descriptor.fields)

    output = tmp_path / "paragraph-fill.docx"
    service.fill_docx(
        template_content=content,
        output_path=output,
        descriptor=descriptor,
        field_values={field["field_key"]: "中国银行自贡分行"},
        sections=[],
    )

    result = Document(output)
    assert result.paragraphs[-1].text == "开户行： 中国银行自贡分行"


def test_repository_tender_uses_actual_template_chapter_not_toc(tmp_path):
    samples = list(Path("database/输入（招标文件）").glob("*.docx"))
    if not samples:
        return
    service = ResponseTemplateService()
    content = samples[0].read_bytes()
    descriptor = service.detect(samples[0].name, content)

    assert descriptor.detected is True
    assert descriptor.start_block is not None
    assert descriptor.start_block > 300
    assert descriptor.marker_text.startswith("第七章")
    required = service.required_fields(descriptor.snapshot())
    assert {"project_name", "project_number", "bidder_name"} <= set(required)
    assert all(
        not (
            item["field_key"] == "legal_representative"
            and any(token in item["label"] for token in ("签字", "盖章", "复印件"))
        )
        for item in descriptor.fields
    )
    assert len(required) >= 10
    assert descriptor.end_block is not None
    assert descriptor.end_block > descriptor.start_block

    report = service.fill_docx(
        template_content=content,
        output_path=tmp_path / "real-template-fill.docx",
        descriptor=descriptor,
        field_values={"project_name": "当前项目"},
        sections=[{"title": "技术方案", "content": "受控测试正文。"}],
    )
    assert report.unresolved_sections == ()
    result = Document(tmp_path / "real-template-fill.docx")
    implementation = next(
        item for item in result.paragraphs if "项目实施方案" in item.text
    )
    assert implementation._p.getnext() is not None
    assert "受控测试正文" in "".join(
        node.text or "" for node in implementation._p.getnext().iter()
    )


def test_strict_fill_keeps_source_directory_position_and_refreshes_it(tmp_path):
    document = Document()
    document.add_heading("附件：投标文件格式", level=1)
    document.add_paragraph("目录")
    document.add_heading("一、项目理解", level=1)
    document.add_heading("二、实施计划", level=1)
    stream = BytesIO()
    document.save(stream)

    service = ResponseTemplateService()
    descriptor = service.detect("投标文件格式.docx", stream.getvalue())
    output = tmp_path / "strict-directory.docx"
    service.fill_docx(
        template_content=stream.getvalue(),
        output_path=output,
        descriptor=descriptor,
        field_values={},
        sections=[
            {"title": "一、项目理解", "content": "项目理解正文。"},
            {"title": "二、实施计划", "content": "实施计划正文。"},
        ],
    )

    result = Document(output)
    paragraphs = [item.text for item in result.paragraphs]
    assert paragraphs.index("目录") < paragraphs.index("一、项目理解")
    assert paragraphs.index("一、项目理解") < paragraphs.index("项目理解正文。")
    assert paragraphs.index("二、实施计划") < paragraphs.index("实施计划正文。")
    assert result.settings.element.find(qn("w:updateFields")).get(
        qn("w:val")
    ) == "true"
