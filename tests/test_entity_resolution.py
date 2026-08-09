from io import BytesIO
from uuid import uuid4

from docx import Document

from app.core.entity_resolution import (
    EntityCandidate,
    EntityType,
    EntityResolutionContext,
    EntityResolutionEngine,
    FillStrategy,
    Organization,
    Person,
    ProjectRole,
    ProjectRoleAssignment,
    SlotContextClassifier,
)
from app.services.generation_profile_service import (
    GenerationProfile,
    GenerationProfileService,
)
from app.services.response_template_service import ResponseTemplateService


def _authorization_template() -> bytes:
    document = Document()
    document.add_heading("附件：投标文件格式", level=1)
    document.add_heading("法定代表人身份证明", level=2)
    document.add_paragraph(
        "本人___（姓名）系___（投标人名称）的法定代表人。"
    )
    document.add_heading("授权委托书", level=2)
    document.add_paragraph("现委托___（姓名）为我方代理人。")
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _slot(label: str, context: str):
    return SlotContextClassifier.classify(
        label=label,
        surrounding_text=context,
        source_location="授权委托书/正文第1段",
        document_section="授权委托书",
    )


def test_authorization_template_resolves_organization_and_distinct_roles():
    descriptor = ResponseTemplateService().detect(
        "采购文件.docx", _authorization_template()
    )

    assert descriptor.detected is True
    slots = [item for item in descriptor.fields if item["semantic_field"]]
    assert any(
        item["semantic_field"] == "organization.full_name"
        and item["expected_entity_type"] == "Organization"
        for item in slots
    )
    assert any(
        item["semantic_field"] == "person.name"
        and item["expected_role"] == "LEGAL_REPRESENTATIVE"
        for item in slots
    )
    assert any(
        item["semantic_field"] == "person.name"
        and item["expected_role"] == "AUTHORIZED_REPRESENTATIVE"
        for item in slots
    )
    assert all(item["surrounding_text"] for item in slots)
    assert all(item["paragraph_index"] for item in slots)
    person_slots = [
        item for item in slots if item["semantic_field"] == "person.name"
    ]
    assert [item["expected_role"] for item in person_slots] == [
        "LEGAL_REPRESENTATIVE", "AUTHORIZED_REPRESENTATIVE",
    ]
    organization_slot = next(
        item for item in slots
        if item["semantic_field"] == "organization.full_name"
    )
    assert organization_slot["expected_role"] is None
    assert organization_slot["relation_path"] == [
        "当前项目", "投标人", "企业全称",
    ]


def test_mixed_authorization_sentence_binds_blank_to_nearest_role():
    slot = _slot(
        "姓名",
        "投标人的法定代表人，现委托_______（姓名）为我方代理人。",
    )

    assert slot.expected_role is ProjectRole.AUTHORIZED_REPRESENTATIVE
    assert slot.relation_path == (
        "当前项目", "授权代表", "姓名",
    )


def test_legal_representative_aliases_are_one_role():
    for alias in ("法人", "法定代表", "法定代表人"):
        slot = _slot("姓名", f"{alias}姓名：___")
        assert slot.expected_role is ProjectRole.LEGAL_REPRESENTATIVE


def test_multiple_people_are_candidates_but_never_randomly_selected():
    project_id = uuid4()
    organization_id = uuid4()
    people = (
        Person(id=uuid4(), name="人员甲"),
        Person(id=uuid4(), name="人员乙"),
    )
    candidates = tuple(
        EntityCandidate(
            person_id=person.id,
            name=person.name,
            title=None,
            match_basis="同一投标主体的已核验人员",
            source_document="人员库",
            source_location="人员档案",
            confidence=1.0,
        )
        for person in people
    )
    context = EntityResolutionContext(
        project_id=project_id,
        organization=Organization(id=organization_id, full_name="测试公司"),
        people=people,
        candidates_by_role={
            ProjectRole.AUTHORIZED_REPRESENTATIVE: candidates
        },
    )

    result = EntityResolutionEngine().resolve(
        _slot("姓名", "现委托___（姓名）为我方代理人"), context
    )

    assert result.status == "binding_required"
    assert result.person is None
    assert len(result.candidates) == 2
    assert "不会随机选择" in result.reason


def test_role_binding_keeps_all_person_attributes_on_one_person_id():
    project_id = uuid4()
    organization_id = uuid4()
    authorized = Person(
        id=uuid4(),
        name="侯明",
        title="项目总监",
        id_number="110101199001011234",
        certificates=({"type": "咨询工程师证书", "attachment_id": "A1"},),
        source_documents=({
            "title": "人员资料库",
            "location": "授权人员甲档案",
            "excerpt": "姓名、职务和证件经人工核验",
        },),
    )
    other = Person(id=uuid4(), name="其他人员", title="经理")
    assignment = ProjectRoleAssignment(
        project_id=project_id,
        role=ProjectRole.AUTHORIZED_REPRESENTATIVE,
        person_id=authorized.id,
        organization_id=organization_id,
    )
    context = EntityResolutionContext(
        project_id=project_id,
        project_name="测试项目",
        organization=Organization(id=organization_id, full_name="测试公司"),
        people=(authorized, other),
        assignments=(assignment,),
    )
    fields = []
    for key, label in (
        ("authorized_representative", "授权代表姓名"),
        ("person_title", "授权代表职务"),
        ("person_id_number", "授权代表身份证号码"),
    ):
        slot = SlotContextClassifier.classify(
            label=label,
            surrounding_text=f"授权委托书 {label}：___",
            source_location=f"授权委托书/{label}",
            document_section="授权委托书",
            canonical_hint=key,
        )
        fields.append({"field_key": key, "label": label, **slot.snapshot()})
    profile = GenerationProfile(
        project_id=project_id,
        generation_mode="strict_template",
        historical_case_mode="closest_case",
        template_descriptor={"fields": fields},
        template_field_values={},
    )

    decisions = {
        item["field_key"]: item
        for item in GenerationProfileService.template_field_decisions(
            profile, entity_context=context
        )
    }

    assert decisions["authorized_representative"]["value"] == authorized.name
    assert decisions["person_title"]["value"] == authorized.title
    assert decisions["person_id_number"]["value"] == authorized.id_number
    assert all(item["binding_status"] == "resolved" for item in decisions.values())
    assert authorized.certificates[0]["attachment_id"] == "A1"
    assert all("同一 person_id" in " ".join(item["match_path"]) for item in decisions.values())


def test_ambiguous_legal_or_authorized_representative_needs_review():
    slot = _slot("姓名", "法人或授权代表（姓名）：___")

    assert slot.expected_entity_type.value == "Person"
    assert slot.expected_role is None


def test_date_and_seal_are_document_relationships_not_isolated_values():
    document = Document()
    document.add_heading("附件：投标文件格式", level=1)
    document.add_paragraph("投标人名称（加盖公章）：________")
    document.add_paragraph("盖章：________")
    document.add_paragraph("日 期：________")
    stream = BytesIO()
    document.save(stream)

    descriptor = ResponseTemplateService().detect(
        "采购文件.docx", stream.getvalue()
    )

    assert [item["display_name"] for item in descriptor.fields] == [
        "当前项目投标人名称", "本项目投标文件签署日期",
    ]
    bidder = descriptor.fields[0]
    assert bidder["required_actions"] == ["加盖投标人公章"]
    signing_date = descriptor.fields[1]
    assert signing_date["semantic_field"] == "bid_response.signing_date"
    assert signing_date["relation_path"] == [
        "当前项目", "投标文件", "签署日期",
    ]
    assert len(descriptor.actions) == 1
    assert descriptor.actions[0]["display_name"] == "加盖投标人公章"


def test_two_party_seal_block_is_actions_not_an_empty_text_field():
    document = Document()
    document.add_heading("附件：投标文件格式", level=1)
    document.add_paragraph("甲方（盖章）：_________ 乙方（盖章）：_________")
    stream = BytesIO()
    document.save(stream)

    descriptor = ResponseTemplateService().detect(
        "采购文件.docx", stream.getvalue()
    )

    assert descriptor.fields == ()
    assert {item["display_name"] for item in descriptor.actions} == {
        "甲方签章", "乙方签章",
    }


def test_project_name_and_number_in_one_slot_is_a_composed_relation():
    slot = _slot(
        "项目名称，项目编号/包号",
        "我方参加___（项目名称，项目编号/包号）组织的招标活动。",
    )

    assert slot.canonical_key == "project_reference"
    assert slot.semantic_field == "project.reference"
    assert slot.display_name == "当前项目名称及编号"
    assert slot.fill_strategy.value == "composed_value"
    assert slot.relation_path == ("当前项目", "组合项目名称与项目编号")


def test_unknown_slot_never_exposes_custom_internal_identifier():
    slot = _slot("其他说明", "其他说明：________")

    assert slot.canonical_key == "unmapped_field"
    assert "custom_" not in slot.canonical_key
    assert slot.display_name == "尚未识别的业务槽位"


def test_table_columns_bind_to_business_objects_without_guessing_a_person():
    person = SlotContextClassifier.classify(
        label="证书名称",
        surrounding_text="项目团队人员表 | 【当前空位：证书名称】",
        source_location="表格4/第3行/第6列",
        document_section="响应文件格式 / 项目团队表",
    )
    organization = SlotContextClassifier.classify(
        label="营业执照号",
        surrounding_text="供应商基本情况 | 【当前空位：营业执照号】",
        source_location="表格1/第10行/第2列",
        document_section="响应文件格式 / 供应商基本情况表",
    )

    assert person.expected_entity_type is EntityType.PERSON
    assert person.expected_role is None
    assert person.fill_strategy is FillStrategy.UNRESOLVED
    assert person.ontology_concept == "Person[UNBOUND_ROW].certificate_name"
    assert organization.expected_entity_type is EntityType.ORGANIZATION
    assert organization.ontology_concept.endswith(".business_license_number")


def test_role_attribute_and_staffing_total_are_not_misread_as_person_names():
    professional_title = SlotContextClassifier.classify(
        label="技术职称",
        surrounding_text=(
            "法定代表人 | 姓名 | 技术职称 | 联系电话 | "
            "【当前空位：技术职称】"
        ),
        source_location="表格1/第6行/第6列",
        document_section="响应文件格式 / 供应商基本情况表",
    )
    project_manager_total = SlotContextClassifier.classify(
        label="项目经理",
        surrounding_text=(
            "员工总人数 | 其中 | 高级职称人员 | 项目经理 | "
            "【当前空位：项目经理】"
        ),
        source_location="表格1/第9行/第10列",
        document_section="响应文件格式 / 供应商基本情况表",
    )

    assert professional_title.expected_role is ProjectRole.LEGAL_REPRESENTATIVE
    assert professional_title.semantic_field == "person.professional_title"
    assert professional_title.canonical_key == "person_professional_title"
    assert project_manager_total.expected_role is None
    assert project_manager_total.expected_entity_type is EntityType.ORGANIZATION
    assert project_manager_total.semantic_field == "organization.project_manager_count"


def test_structured_word_tables_bind_cells_to_upper_level_business_objects():
    cases = SlotContextClassifier.classify(
        label="用户/业主名称",
        surrounding_text=(
            "序号 | 用户/业主名称 | 项目名称 | 项目内容 | 合同总价 | "
            "签订时间 | 完成时间 | 【当前空位：用户/业主名称】"
        ),
        source_location="表格1/第2行/第2列",
        document_section="响应文件格式 / 响应供应商同类项目经验",
        table_index=1,
        row=2,
        column=2,
    )
    response = SlotContextClassifier.classify(
        label="响应供应商响应描述",
        surrounding_text=(
            "序号 | 磋商文件条款描述 | 响应供应商响应描述 | "
            "偏离情况说明 | 【当前空位：响应供应商响应描述】"
        ),
        source_location="表格2/第2行/第3列",
        document_section="响应文件格式 / 采购需求响应一览表",
        table_index=2,
        row=2,
        column=3,
    )
    certificate = SlotContextClassifier.classify(
        label="发证单位",
        surrounding_text=(
            "证书名称 | 发证单位 | 证书等级 | 证书有效期 | "
            "【当前空位：发证单位】"
        ),
        source_location="表格3/第2行/第2列",
        document_section="响应文件格式 / 响应供应商证书一览表",
        table_index=3,
        row=2,
        column=2,
    )
    sequence = SlotContextClassifier.classify(
        label="序号",
        surrounding_text="序号 | 项目名称 | 【当前空位：序号】",
        source_location="表格1/第2行/第1列",
        document_section="响应文件格式",
        table_index=1,
        row=2,
        column=1,
    )
    service_term = SlotContextClassifier.classify(
        label="服务期限",
        surrounding_text=(
            "采购内容 | 数量 | 磋商报价 | 服务期限 | "
            "【当前空位：服务期限】"
        ),
        source_location="表格4/第2行/第4列",
        document_section="响应文件格式 / 磋商报价表",
        table_index=4,
        row=2,
        column=4,
    )

    assert cases.expected_entity_type is EntityType.BUSINESS_CASE
    assert cases.semantic_field == "business_case.client_name"
    assert cases.fill_strategy is FillStrategy.KNOWLEDGE_COLLECTION
    assert response.expected_entity_type is EntityType.RESPONSE_ITEM
    assert response.semantic_field == "bid_response.response_item.response_text"
    assert response.fill_strategy is FillStrategy.GENERATED_COLLECTION
    assert certificate.expected_entity_type is EntityType.CERTIFICATE
    assert certificate.semantic_field == "certificate.issuer"
    assert sequence.expected_entity_type is None
    assert sequence.fill_strategy is FillStrategy.AUTO_LAYOUT
    assert service_term.expected_entity_type is EntityType.PROJECT
    assert service_term.semantic_field == "project.service_term"


def test_standalone_contact_phone_binds_to_project_contact_person():
    slot = SlotContextClassifier.classify(
        label="电话",
        surrounding_text="【当前空位：电话】",
        source_location="第111段/第1个字段",
        document_section="第六章 响应文件格式",
        paragraph_index=111,
        canonical_hint="contact_phone",
    )

    assert slot.semantic_field == "person.phone"
    assert slot.expected_role is ProjectRole.CONTACT_PERSON
    assert slot.display_name == "联系人联系电话"


def test_table_collections_are_reviewed_as_business_objects_not_every_cell():
    document = Document()
    document.sections[0].header.paragraphs[0].text = "原模板页眉"
    document.sections[0].footer.paragraphs[0].text = "原模板页脚"
    document.add_heading("第六章 响应文件格式", level=1)

    cases = document.add_table(rows=3, cols=7)
    for index, label in enumerate((
        "序号", "用户/业主名称", "项目名称", "项目内容",
        "合同总价", "签订时间", "完成时间",
    )):
        cases.cell(0, index).text = label

    responses = document.add_table(rows=3, cols=4)
    for index, label in enumerate((
        "序号", "磋商文件条款描述", "响应供应商响应描述", "偏离情况说明",
    )):
        responses.cell(0, index).text = label

    stream = BytesIO()
    document.save(stream)
    descriptor = ResponseTemplateService().detect(
        "采购文件.docx", stream.getvalue()
    )
    profile = GenerationProfile(
        project_id=uuid4(),
        generation_mode="strict_template",
        historical_case_mode="closest_case",
        template_descriptor=descriptor.snapshot(),
        template_field_values={},
    )
    variables = GenerationProfileService.template_variable_decisions(profile)

    assert descriptor.fields
    assert all(item["semantic_field"] != "text.value" for item in descriptor.fields)
    assert all("页眉" not in item["surrounding_text"] for item in descriptor.fields)
    assert all("页脚" not in item["surrounding_text"] for item in descriptor.fields)
    layout = [item for item in descriptor.fields if item["fill_strategy"] == "auto_layout"]
    assert layout and all(item["required"] is False for item in layout)
    review_groups = {
        item["review_group_key"]: item["review_group_label"]
        for item in variables
        if item["resolution_state"] != "layout_managed"
    }
    assert set(review_groups.values()) == {
        "企业业绩表（系统整表匹配）",
        "招标响应表（系统整表生成）",
    }
