import importlib.util
import json
from io import BytesIO
from pathlib import Path

from docx import Document


SCRIPT = (
    Path(__file__).parents[1]
    / "scripts"
    / "audit_template_case_readiness.py"
)
SPEC = importlib.util.spec_from_file_location("readiness_audit", SCRIPT)
MODULE = importlib.util.module_from_spec(SPEC)
assert SPEC.loader is not None
SPEC.loader.exec_module(MODULE)


def _docx(*values: str, headings: bool = False) -> bytes:
    document = Document()
    for value in values:
        if headings:
            document.add_heading(value, level=1)
        else:
            document.add_paragraph(value)
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _manifest(tmp_path, count: int):
    pairs = []
    for index in range(count):
        case = tmp_path / f"case-{index}"
        case.mkdir()
        (case / "tender.docx").write_bytes(
            _docx("采购需求", "服务方案要求")
        )
        (case / "winning.docx").write_bytes(
            _docx("第一章 项目理解", "1.1 总体思路", headings=True)
        )
        pairs.append({
            "tender": f"case-{index}/tender.docx",
            "winning_proposal": f"case-{index}/winning.docx",
            "project_type": "咨询服务",
            "industry": "政府咨询",
            "quality_score": 0.9,
        })
    path = tmp_path / "pairs.json"
    path.write_text(json.dumps({"pairs": pairs}), encoding="utf-8")
    return path


def test_readiness_audit_reports_missing_real_pairs_without_writing(tmp_path):
    report = MODULE.audit(_manifest(tmp_path, 1), expected_pairs=5)

    assert report["validated_pairs"] == 1
    assert report["ready"] is False
    assert report["blocker"] == "仍缺 4 组真实案例文件。"
    assert report["privacy"]["database_write"] is False


def test_readiness_audit_accepts_complete_isolated_five_pair_batch(tmp_path):
    report = MODULE.audit(_manifest(tmp_path, 5), expected_pairs=5)

    assert report["ready"] is True
    assert all(report["gates"].values())
    assert report["generation_modes"] == {"planned": 5}
