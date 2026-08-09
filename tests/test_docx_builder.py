from docx import Document

from app.services.docx_builder import (
    _clean_export_markdown,
    _strip_leading_heading,
    build_full_proposal_docx,
    build_proposal_docx,
    delivery_title,
)


def test_build_proposal_docx_with_response_matrix(tmp_path):
    output = tmp_path / "proposal.docx"

    build_proposal_docx(
        output,
        project_name="测试项目",
        section_title="实施方案",
        content="# 总体思路\n\n系统采用分阶段实施。\n\n- 每日备份",
        requirements=[
            {
                "normalized_text": "系统须支持每日备份。",
                "quote": "系统须支持每日备份。",
                "sources": [
                    {
                        "filename": "招标文件.pdf",
                        "locator": {
                            "kind": "page",
                            "page": 12,
                            "paragraph_start": None,
                            "paragraph_end": None,
                        },
                    }
                ],
            }
        ],
    )

    assert output.is_file()
    document = Document(output)
    text = "\n".join(item.text for item in document.paragraphs)
    assert "实施方案" in text
    assert "总体思路" in text
    assert len(document.tables) == 1
    assert document.tables[0].cell(1, 2).text.startswith(
        "本项目招标文件，第 12 页"
    )


def test_export_markdown_removes_internal_prompt_labels():
    content = (
        "依据 Requirement: 123e4567-e89b-12d3-a456-426614174000 "
        "并参考 Matched Knowledge "
        "123e4567-e89b-12d3-a456-426614174001。"
    )

    cleaned = _clean_export_markdown(content)

    assert "Requirement" not in cleaned
    assert "Matched Knowledge" not in cleaned
    assert "f47ac10b-58cc-4372-a567-0e02b2c3d479" not in cleaned
    assert "123e4567" not in cleaned


def test_duplicate_leading_chapter_heading_is_removed():
    assert _strip_leading_heading(
        "# 服务范围与工作内容\n\n正文",
        "服务范围与工作内容",
    ).strip() == "正文"


def test_full_proposal_uses_delivery_title_and_directory(tmp_path):
    output = tmp_path / "full.docx"
    build_full_proposal_docx(
        output,
        project_name="测试项目",
        sections=[{"title": "实施方案", "content": "正文"}],
        requirements=[],
    )

    document = Document(output)
    text = "\n".join(item.text for item in document.paragraphs)
    assert delivery_title("测试项目") == "《AI投标文件+测试项目》"
    assert "《AI投标文件+测试项目》" in text
    assert "目录" in text
    assert "1. 实施方案" in text
    assert document.core_properties.title == "《AI投标文件+测试项目》"
