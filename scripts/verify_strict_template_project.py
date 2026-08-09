from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from uuid import UUID

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from app.knowledge.case_fact_resolver import CaseFactResolver
from app.knowledge.enterprise_fact_resolver import EnterpriseFactResolver
from app.services.export_service import ExportService
from app.services.generation_profile_service import GenerationProfileService
from app.services.section_service import SectionService


def main() -> int:
    parser = argparse.ArgumentParser(description="验收严格模板回填项目。")
    parser.add_argument("project_id", type=UUID)
    parser.add_argument("--approve-and-export", action="store_true")
    args = parser.parse_args()
    if not args.approve_and_export:
        parser.error("必须显式传入 --approve-and-export。")

    profile_service = GenerationProfileService()
    profile = profile_service.get(args.project_id)
    if profile.generation_mode != "strict_template":
        raise SystemExit("项目未进入严格模板分支。")
    decisions = profile_service.template_field_decisions(
        profile,
        enterprise_facts=EnterpriseFactResolver().resolve(args.project_id),
        case_candidates=CaseFactResolver().resolve(args.project_id),
    )
    for item in decisions:
        if item["status"] == "REVIEW_REQUIRED" and item["value"]:
            profile_service.review_template_field(
                args.project_id, item["field_key"], "confirm"
            )

    section_service = SectionService()
    for section in section_service.list(args.project_id):
        if not section.get("current_version"):
            raise SystemExit(f"章节未生成：{section['title']}")
        if section.get("status") != "approved":
            section_service.approve(args.project_id, section["id"])

    exported = ExportService().create_full(args.project_id)
    path, _ = ExportService().download_info(args.project_id, exported["id"])
    document = Document(path)
    text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    report = {
        "status": exported["status"],
        "filename": exported["filename"],
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "nonempty_chars": len(re.sub(r"\s+", "", text)),
        "internal_identifier_leaks": len(re.findall(r"custom_[0-9a-f]{8,}", text)),
        "placeholder_x_runs": len(re.findall(r"X{3,}|x{3,}", text)),
        "review_markers": text.count("【待审核"),
        "review_candidate_count": sum(item["status"] == "REVIEW_REQUIRED" for item in decisions),
        "missing_field_count": sum(item["status"] == "MISSING" for item in decisions),
    }
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if report["internal_identifier_leaks"] == 0 and report["placeholder_x_runs"] == 0 else 1


if __name__ == "__main__":
    raise SystemExit(main())
