from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from io import BytesIO
from statistics import median
from typing import TYPE_CHECKING, Any
from uuid import UUID

from docx import Document

from app.knowledge.permissions import KnowledgeAccessContext

if TYPE_CHECKING:
    from app.memory.engine import ProposalMemoryEngine


@dataclass(frozen=True)
class HistoricalCasePair:
    tender_filename: str
    tender_text: str
    proposal_filename: str
    proposal_content: bytes
    project_type: str
    industry: str
    quality_score: float = 0.85


class HistoricalCasePatternExtractor:
    """Extracts reusable form, never historical prose or enterprise facts."""

    HEADING_RE = re.compile(
        r"^(?:第?[一二三四五六七八九十百\d]+[章节、.]|"
        r"[一二三四五六七八九十]+、|\d+(?:\.\d+){0,3})\s*(.+)$"
    )

    def extract(self, content: bytes) -> list[dict[str, Any]]:
        document = Document(BytesIO(content))
        paragraphs = [item for item in document.paragraphs if item.text.strip()]
        headings: list[tuple[int, str, int]] = []
        for index, paragraph in enumerate(paragraphs):
            text = paragraph.text.strip()
            style_name = paragraph.style.name if paragraph.style else ""
            match = self.HEADING_RE.match(text)
            if style_name.startswith("Heading") or match:
                level = self._heading_level(style_name, text)
                title = self._generic_heading(match.group(1) if match else text)
                headings.append((index, title, level))
        if not headings:
            headings = [(0, "方案正文", 1)]

        patterns: list[dict[str, Any]] = []
        for position, (start, title, level) in enumerate(headings):
            boundary = next(
                (
                    candidate
                    for candidate in range(position + 1, len(headings))
                    if headings[candidate][2] <= level
                ),
                len(headings),
            )
            end = (
                headings[boundary][0]
                if boundary < len(headings)
                else len(paragraphs)
            )
            body = paragraphs[start + 1 : end]
            lengths = [len(item.text.strip()) for item in body]
            child_titles = [
                item[1] for item in headings[position + 1 : boundary]
                if item[2] > level
            ][:8]
            table_shapes = [
                {
                    "columns": len(table.columns),
                    "header_roles": self._generic_headers(table),
                }
                for table in document.tables
            ][:8]
            patterns.append(
                {
                    "chapter_structure": [title, *child_titles],
                    "analysis_dimensions": self._dimensions(title, child_titles),
                    "writing_method": {
                        "paragraph_count": len(body),
                        "median_paragraph_length": int(median(lengths)) if lengths else 0,
                        "uses_tables": bool(table_shapes),
                        "table_patterns": table_shapes,
                    },
                    "visual_pattern": {
                        "heading_level": level,
                        "table_count": len(document.tables),
                    },
                    "prohibited_fact_copy": True,
                    "source_facts_removed": True,
                }
            )
        return patterns

    @staticmethod
    def _heading_level(style_name: str, text: str) -> int:
        match = re.search(r"(\d+)$", style_name)
        if match:
            return max(1, min(3, int(match.group(1))))
        if re.match(r"^第?[一二三四五六七八九十百\d]+章", text):
            return 1
        if re.match(r"^[一二三四五六七八九十]+、", text):
            return 2
        return 3

    @staticmethod
    def _generic_heading(value: str) -> str:
        value = re.sub(r"[\s:：]+", "", value)
        known = (
            "项目理解", "总体思路", "技术方案", "实施计划", "进度安排",
            "组织管理", "人员配置", "质量保障", "验收方案", "培训方案",
            "运维服务", "安全方案", "应急预案", "服务承诺",
        )
        return next((item for item in known if item in value), "通用响应章节")

    @staticmethod
    def _dimensions(title: str, children: list[str]) -> list[str]:
        dimensions = list(dict.fromkeys([title, *children]))
        return [item for item in dimensions if item != "通用响应章节"][:8]

    @staticmethod
    def _generic_headers(table) -> list[str]:
        if not table.rows:
            return []
        roles = []
        for cell in table.rows[0].cells:
            text = cell.text.strip()
            role = next(
                (
                    item for item in (
                        "序号", "事项", "要求", "响应", "偏离", "人员",
                        "职责", "阶段", "时间", "成果", "风险", "措施",
                    )
                    if item in text
                ),
                "通用字段",
            )
            roles.append(role)
        return roles


class HistoricalCaseLearningService:
    def __init__(
        self,
        memory_engine: "ProposalMemoryEngine | None" = None,
        extractor: HistoricalCasePatternExtractor | None = None,
    ):
        if memory_engine is None:
            from app.memory.engine import ProposalMemoryEngine

            memory_engine = ProposalMemoryEngine()
        self.memory_engine = memory_engine
        self.extractor = extractor or HistoricalCasePatternExtractor()

    def learn_pairs(
        self,
        *,
        access_context: KnowledgeAccessContext,
        pairs: list[HistoricalCasePair],
    ) -> list[UUID]:
        learned: list[UUID] = []
        for pair in pairs:
            if not pair.tender_text.strip():
                raise ValueError("招标文件未解析出可用文本。")
            if not pair.proposal_content:
                raise ValueError("中标响应文件为空。")
            if not pair.project_type.strip() or not pair.industry.strip():
                raise ValueError("项目类型和行业不能为空。")
            digest = hashlib.sha256(
                pair.tender_text.encode("utf-8") + pair.proposal_content
            ).hexdigest()
            for chapter_index, pattern in enumerate(
                self.extractor.extract(pair.proposal_content), start=1
            ):
                pattern["source_pair_checksum"] = digest
                pattern["reference_scope"] = "organization_private"
                pattern["chapter_index"] = chapter_index
                learned.append(
                    self.memory_engine.add_pattern(
                        access_context=access_context,
                        project_type=pair.project_type,
                        industry=pair.industry,
                        chapter_title=pattern["chapter_structure"][0],
                        pattern=pattern,
                        quality_score=pair.quality_score,
                    )
                )
        return learned
