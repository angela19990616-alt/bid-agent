import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 100, 110)
TABLE_FILL = "F4F6F9"


def build_proposal_docx(
    output_path: Path,
    *,
    project_name: str,
    section_title: str,
    content: str,
    requirements: list[dict],
) -> None:
    document = Document()
    _configure_document(document, project_name, section_title)
    _add_title_block(document, project_name, section_title)
    _add_markdown(document, content)
    document.add_page_break()
    document.add_heading("要求响应与来源清单", level=1)
    note = document.add_paragraph(
        "以下清单用于人工复核章节响应范围，原文引用以系统保存的来源定位为准。"
    )
    note.style = document.styles["Normal"]
    _add_requirement_table(document, requirements)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _configure_document(
    document: Document,
    project_name: str,
    section_title: str,
) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Arial Unicode MS"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style._element.rPr.rFonts.set(
            qn("w:eastAsia"),
            "Arial Unicode MS",
        )
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    document.core_properties.author = "Bid Agent"
    document.core_properties.last_modified_by = "Bid Agent"
    document.core_properties.title = section_title
    document.core_properties.subject = project_name

    header = section.header.paragraphs[0]
    header.text = project_name
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _style_runs(header, 9, MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("第 ")
    _set_run_font(run, 9, MUTED)
    _append_page_field(footer)
    run = footer.add_run(" 页")
    _set_run_font(run, 9, MUTED)


def _add_title_block(
    document: Document,
    project_name: str,
    section_title: str,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(36)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(project_name)
    _set_run_font(run, 12, MUTED, bold=True)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(section_title)
    _set_run_font(run, 24, RGBColor(0, 0, 0), bold=True)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(24)
    run = paragraph.add_run("技术方案章节")
    _set_run_font(run, 14, MUTED)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(28)
    run = paragraph.add_run(
        f"导出时间：{datetime.now().astimezone():%Y-%m-%d %H:%M}"
    )
    _set_run_font(run, 10, MUTED)


def _add_markdown(document: Document, content: str) -> None:
    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            document.add_heading(heading.group(2), level=len(heading.group(1)))
            continue
        if re.match(r"^[-*]\s+", line):
            paragraph = document.add_paragraph(
                re.sub(r"^[-*]\s+", "", line),
                style="List Bullet",
            )
            _configure_list_paragraph(paragraph)
            continue
        if re.match(r"^\d+[.)、]\s*", line):
            paragraph = document.add_paragraph(
                re.sub(r"^\d+[.)、]\s*", "", line),
                style="List Number",
            )
            _configure_list_paragraph(paragraph)
            continue
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_inline_text(paragraph, line)


def _add_inline_text(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        value = part[2:-2] if bold else part
        run = paragraph.add_run(value)
        _set_run_font(run, 11, RGBColor(0, 0, 0), bold=bold)


def _configure_list_paragraph(paragraph) -> None:
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.194)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.208
    _style_runs(paragraph, 11, RGBColor(0, 0, 0))


def _add_requirement_table(document: Document, requirements: list[dict]) -> None:
    table = document.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [1440, 3960, 3960]
    headers = ("序号", "确认要求", "原文来源")
    for index, (cell, label) in enumerate(zip(table.rows[0].cells, headers)):
        cell.text = label
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.width = Inches(widths[index] / 1440)
        _set_cell_fill(cell, TABLE_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _style_runs(paragraph, 10, DARK_BLUE, bold=True)

    for number, requirement in enumerate(requirements, start=1):
        cells = table.add_row().cells
        source_text = "；".join(
            _source_label(source) for source in requirement["sources"]
        )
        values = (
            str(number),
            requirement["normalized_text"],
            f"{source_text}\n原文：{requirement['quote']}",
        )
        for index, (cell, value) in enumerate(zip(cells, values)):
            cell.text = value
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if index == 0
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            _style_runs(paragraph, 9.5, RGBColor(0, 0, 0))
    _set_table_geometry(table, widths)
    _set_repeat_table_header(table.rows[0])


def _source_label(source: dict) -> str:
    locator = source["locator"]
    if locator["kind"] == "page":
        location = f"第 {locator['page']} 页"
    else:
        start = locator["paragraph_start"]
        end = locator["paragraph_end"]
        location = (
            f"第 {start} 段"
            if start == end
            else f"第 {start}-{end} 段"
        )
    return f"{source['filename']}，{location}"


def _set_table_geometry(table, widths: list[int]) -> None:
    table_element = table._tbl
    properties = table_element.tblPr
    width = properties.first_child_found_in("w:tblW")
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), "9360")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    properties.append(indent)

    grid = table_element.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)

    for row in table.rows:
        for cell, value in zip(row.cells, widths):
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.get_or_add_tcW()
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(value))
            margins = cell_properties.first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                cell_properties.append(margins)
            for side, amount in (
                ("top", 80),
                ("bottom", 80),
                ("start", 120),
                ("end", 120),
            ):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), str(amount))
                node.set(qn("w:type"), "dxa")
                margins.append(node)


def _set_repeat_table_header(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    row_properties.append(repeat)


def _set_cell_fill(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    properties.append(shading)


def _append_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    _set_run_font(run, 9, MUTED)


def _style_runs(
    paragraph,
    size: float,
    color: RGBColor,
    bold: bool | None = None,
) -> None:
    for run in paragraph.runs:
        _set_run_font(run, size, color, bold=bold)


def _set_run_font(
    run,
    size: float,
    color: RGBColor,
    bold: bool | None = None,
) -> None:
    run.font.name = "Arial Unicode MS"
    run._element.get_or_add_rPr().rFonts.set(
        qn("w:ascii"),
        "Arial Unicode MS",
    )
    run._element.get_or_add_rPr().rFonts.set(
        qn("w:hAnsi"),
        "Arial Unicode MS",
    )
    run._element.get_or_add_rPr().rFonts.set(
        qn("w:eastAsia"),
        "Arial Unicode MS",
    )
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
