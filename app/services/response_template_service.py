from __future__ import annotations

import re
from collections import Counter
from copy import deepcopy
from dataclasses import asdict, dataclass
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader

from app.core.strict_fill import StrictFillDecisionEngine
from app.core.entity_resolution import SlotContextClassifier


TEMPLATE_MARKERS = (
    "投标文件格式", "响应文件格式", "报价文件格式", "资格响应格式",
    "技术响应格式", "附件格式", "响应文件组成及格式",
)
NEGATED_TEMPLATE_PREFIX_RE = re.compile(
    r"(?:未提供|没有提供|不提供|未附|不附|不含|不包括|不存在|未设置|"
    r"无(?:统一|指定|固定)?|自行编制).{0,12}$"
)
STRICT_MARKERS = (
    "不得修改", "不得增删", "格式不变", "严格按照", "按附件格式",
    "按本格式", "原格式",
)
FIELD_ALIASES = {
    "project_name": ("项目名称", "采购项目名称", "招标项目名称"),
    "project_number": ("项目编号", "采购编号", "招标编号"),
    "bidder_name": (
        "供应商名称", "投标人名称", "响应人名称", "企业名称", "单位名称",
    ),
    "legal_representative": ("法定代表人", "法定代表", "法人代表"),
    "authorized_representative": (
        "授权代表", "授权代理人", "委托代理人", "被授权人", "受托人",
    ),
    "project_manager_name": ("项目负责人", "项目经理"),
    "technical_lead_name": ("技术负责人", "技术总监"),
    "signatory_name": ("签字人", "签署人"),
    "person_id_number": ("身份证号码", "身份证号"),
    "person_title": ("职务", "职位", "岗位"),
    "date": ("日期", "响应日期", "投标日期"),
    "registered_address": ("注册地址", "供应商地址", "投标人地址"),
    "postal_code": ("邮政编码", "邮编"),
    "contact_person": ("联系人", "项目联系人"),
    "contact_phone": ("联系电话", "手机", "电话"),
    "fax": ("传真",),
    "website": ("网址", "网站"),
    "enterprise_qualification": ("企业资质等级", "企业资质", "资质等级"),
    "bank_account": ("银行账号", "账号"),
    "bid_round": ("报价轮次", "轮次"),
}
PLACEHOLDER_RE = re.compile(
    r"\{\{\s*([\w\u4e00-\u9fff.-]+)\s*\}\}"
    r"|\$\{\s*([\w\u4e00-\u9fff.-]+)\s*\}"
    r"|[\[\uff3b]\s*([\w\u4e00-\u9fff.-]+)\s*[\]\uff3d]"
)
BLANK_TOKEN_RE = re.compile(r"(?:[_＿]{1,}|…{2,}|\.{3,})")
LABELED_BLANK_RE = re.compile(
    r"(?P<label>[\u4e00-\u9fffA-Za-z0-9（）()、/·\s]{1,32}?)[：:]\s*"
    r"(?P<blank>[_＿]{1,}|…{2,}|\.{3,}|(?=$))"
)
PAREN_LABEL_BLANK_RE = re.compile(
    r"[（(](?P<label>[\u4e00-\u9fffA-Za-z0-9、/·\s]{1,24})[）)]\s*"
    r"(?P<blank>[_＿]{1,}|…{2,}|\.{3,})"
)


@dataclass(frozen=True)
class TemplateDescriptor:
    detected: bool
    source_format: str
    template_kind: str
    fidelity: str
    confidence: float
    start_block: int | None
    marker_text: str | None
    table_count: int
    placeholders: tuple[str, ...]
    field_labels: tuple[str, ...]
    strict_reasons: tuple[str, ...]
    outline: tuple[dict[str, Any], ...] = ()
    fields: tuple[dict[str, Any], ...] = ()
    end_block: int | None = None
    font_profile: dict[str, Any] | None = None
    actions: tuple[dict[str, Any], ...] = ()

    def snapshot(self) -> dict[str, Any]:
        return asdict(self)


@dataclass(frozen=True)
class TemplateFillReport:
    filled_fields: tuple[str, ...]
    unresolved_fields: tuple[str, ...]
    inserted_sections: tuple[str, ...]
    unresolved_sections: tuple[str, ...]

    def snapshot(self) -> dict[str, Any]:
        return asdict(self)


class ResponseTemplateService:
    """Detects and fills response templates without rebuilding their layout."""

    def detect(self, filename: str, content: bytes) -> TemplateDescriptor:
        suffix = Path(filename).suffix.lower()
        if suffix == ".docx":
            return self._detect_docx(filename, content)
        if suffix == ".pdf":
            marker_text = None
            try:
                reader = PdfReader(BytesIO(content))
                for page in reader.pages:
                    text = page.extract_text() or ""
                    marker_text = self._positive_template_marker(text)
                    if marker_text:
                        break
            except Exception:
                # The parser/validator owns malformed-file reporting. Template
                # detection remains conservative and never claims coordinates.
                marker_text = None
            detected = bool(marker_text) or bool(
                self._positive_template_marker(filename)
            )
            return TemplateDescriptor(
                detected=detected,
                source_format="pdf",
                template_kind="pdf_reference",
                fidelity="manual_exact_fill_required",
                confidence=(0.8 if marker_text else 0.65 if detected else 0.0),
                start_block=None,
                marker_text=marker_text,
                table_count=0,
                placeholders=(),
                field_labels=(),
                strict_reasons=("PDF版式仅作为严格参照，不能猜测回填坐标。",),
            )
        return TemplateDescriptor(
            False, suffix.lstrip("."), "none", "planned", 0.0,
            None, None, 0, (), (), (),
        )

    @staticmethod
    def required_fields(descriptor: dict[str, Any]) -> list[str]:
        required = list(descriptor.get("placeholders") or ())
        labels = set(descriptor.get("field_labels") or ())
        for key, aliases in FIELD_ALIASES.items():
            if any(alias in labels for alias in aliases):
                required.append(key)
        required.extend(
            item["field_key"]
            for item in descriptor.get("fields") or ()
            if item.get("required", True) and item.get("field_key")
        )
        return list(dict.fromkeys(required))

    def extract_source_fields(
        self,
        filename: str,
        content: bytes,
    ) -> dict[str, str]:
        """Extract procurement facts that are safe to prefill from source."""
        values, _ = self.extract_source_fields_with_evidence(filename, content)
        return values

    def extract_source_fields_with_evidence(
        self,
        filename: str,
        content: bytes,
    ) -> tuple[dict[str, str], dict[str, dict[str, str]]]:
        """Extract procurement facts together with a readable source locator."""
        suffix = Path(filename).suffix.lower()
        segments: list[tuple[str, str]] = []
        try:
            if suffix == ".docx":
                document = Document(BytesIO(content))
                segments.extend(
                    (item.text, f"正文第 {index} 段")
                    for index, item in enumerate(document.paragraphs, start=1)
                    if item.text.strip()
                )
                for table_index, table in enumerate(document.tables, start=1):
                    for row_index, row in enumerate(table.rows, start=1):
                        for cell_index, cell in enumerate(row.cells, start=1):
                            if cell.text.strip():
                                segments.append((
                                    cell.text,
                                    f"表格 {table_index} · 第 {row_index} 行第 {cell_index} 列",
                                ))
            elif suffix == ".pdf":
                segments.extend(
                    (page.extract_text() or "", f"第 {index} 页")
                    for index, page in enumerate(
                        PdfReader(BytesIO(content)).pages, start=1
                    )
                )
        except Exception:
            return {}, {}
        patterns = {
            "project_number": re.compile(
                r"(?:采购项目编号|招标项目编号|项目编号|采购编号|招标编号)"
                r"\s*(?:[（(][^）)\n]{0,12}[）)])?\s*[：:]\s*"
                r"([A-Za-z0-9][A-Za-z0-9._/-]{3,79})"
            ),
            "project_name": re.compile(
                r"(?:采购项目名称|招标项目名称|项目名称)\s*[：:]\s*"
                r"([^\n|]{4,160})"
            ),
        }
        values: dict[str, str] = {}
        evidence: dict[str, dict[str, str]] = {}
        for text, location in segments:
            for key, pattern in patterns.items():
                if key in values:
                    continue
                match = pattern.search(text)
                if match is None:
                    continue
                value = match.group(1).strip().rstrip("。.;；")
                if not StrictFillDecisionEngine.value_matches_field_type(
                    key, value
                ):
                    continue
                values[key] = value
                evidence[key] = {
                    "title": filename,
                    "location": location,
                    "excerpt": re.sub(r"\s+", " ", text).strip()[:500],
                }
        return values, evidence

    def _detect_docx(self, filename: str, content: bytes) -> TemplateDescriptor:
        document = Document(BytesIO(content))
        blocks = list(document.element.body.iterchildren())
        texts = [self._block_text(block) for block in blocks]
        marker_candidates = [
            index for index, text in enumerate(texts)
            if self._positive_template_marker(text)
        ]
        # Tender tables of contents and explanatory paragraphs repeat the
        # title. Prefer the shortest heading-like occurrence, penalising TOC
        # fields and prose, instead of blindly taking the first/last match.
        marker_index = (
            min(
                marker_candidates,
                key=lambda index: self._marker_candidate_rank(
                    texts[index], index
                ),
            )
            if marker_candidates else None
        )
        tables = len(document.tables)
        standalone = bool(self._positive_template_marker(filename))
        detected = marker_index is not None or (standalone and tables > 0)
        start_block = 0 if standalone else marker_index
        end_block = None
        if start_block is not None and not standalone:
            for index in range(start_block + 1, len(texts)):
                compact = texts[index].replace(" ", "")
                if re.match(
                    r"^第[一二三四五六七八九十百零〇\d]+章",
                    compact,
                ):
                    end_block = index
                    break
        candidate_text = "\n".join(
            texts[start_block:end_block] if start_block is not None else []
        )
        placeholders = tuple(sorted({
            next(group for group in match.groups() if group)
            for match in PLACEHOLDER_RE.finditer(candidate_text)
        }))
        labels = self._fillable_labels(document)
        fields, actions = self._extract_fillable_fields(
            document, start_block, end_block
        )
        strict = tuple(marker for marker in STRICT_MARKERS if marker in candidate_text)
        outline = self._extract_outline(document, start_block)
        font_profile = self._extract_font_profile(document, start_block, end_block)
        confidence = 0.0
        if detected:
            confidence = min(
                0.99,
                0.55 + (0.15 if tables else 0) +
                (0.15 if placeholders or labels else 0) +
                (0.1 if strict else 0),
            )
        return TemplateDescriptor(
            detected=detected,
            source_format="docx",
            template_kind=(
                "standalone_attachment" if standalone
                else "embedded_response_section" if detected else "none"
            ),
            fidelity="exact_template" if detected else "planned",
            confidence=round(confidence, 3),
            start_block=start_block,
            marker_text=(texts[marker_index] if marker_index is not None else None),
            table_count=tables,
            placeholders=placeholders,
            field_labels=labels,
            strict_reasons=strict,
            outline=outline,
            fields=fields,
            end_block=end_block,
            font_profile=font_profile,
            actions=actions,
        )

    @staticmethod
    def _positive_template_marker(text: str) -> str | None:
        """Return a format marker only when it is asserted, not negated.

        Tender documents sometimes explicitly say that no response template is
        provided.  Keyword-only detection would route those documents into the
        strict-template writer and block the correct planned-proposal path.
        """
        for marker in TEMPLATE_MARKERS:
            start = 0
            while True:
                index = text.find(marker, start)
                if index < 0:
                    break
                prefix = text[max(0, index - 24):index]
                if not NEGATED_TEMPLATE_PREFIX_RE.search(prefix):
                    return marker
                start = index + len(marker)
        return None

    @staticmethod
    def _marker_candidate_rank(text: str, index: int) -> tuple[int, int, int, int]:
        """Prefer the real response chapter over TOC/prose repetitions.

        Procurement documents commonly repeat the response-format title in
        the table of contents, preparation instructions and explanatory
        notes.  A later formal chapter heading is a stronger boundary signal
        than a slightly shorter earlier spelling (for example, one versus two
        spaces between the chapter number and title).
        """
        compact = re.sub(r"\s+", "", text)
        is_chapter_heading = bool(re.match(
            r"^第[一二三四五六七八九十百零〇\d]+章(?:投标|响应|报价|资格|技术|附件)",
            compact,
        ))
        is_explanatory = bool(re.match(
            r"^(?:说明|备注|注)[：:]?", compact
        )) or "不属于响应文件格式" in compact
        is_toc = "HYPERLINK" in text or "PAGEREF" in text
        return (
            0 if is_chapter_heading else 1,
            1 if is_toc or is_explanatory else 0,
            -index,
            len(text),
        )

    @classmethod
    def _extract_font_profile(
        cls,
        document: DocumentType,
        start_block: int | None,
        end_block: int | None,
    ) -> dict[str, Any]:
        """Record template typography so generated runs inherit, not override, it."""
        body_counter: Counter[tuple] = Counter()
        heading_counter: Counter[tuple] = Counter()
        detected_fonts: set[str] = set()
        for paragraph in cls._paragraphs_in_block_range(
            document, start_block, end_block
        ):
            text = paragraph.text.strip()
            if not text:
                continue
            level, _ = cls._heading_level(paragraph, text)
            for run in paragraph.runs:
                if not run.text.strip():
                    continue
                style = cls._run_style_snapshot(run, paragraph)
                for value in (
                    style.get("ascii"), style.get("hAnsi"),
                    style.get("eastAsia"), style.get("cs"),
                ):
                    if value:
                        detected_fonts.add(value)
                key = tuple(sorted(style.items()))
                (heading_counter if level else body_counter)[key] += len(run.text)

        def most_common(counter: Counter[tuple]) -> dict[str, Any]:
            return dict(counter.most_common(1)[0][0]) if counter else {}

        body = most_common(body_counter)
        heading = most_common(heading_counter) or body
        return {
            "detected_fonts": sorted(detected_fonts),
            "body": body,
            "heading": heading,
            "policy": "inherit_source_template",
        }

    @staticmethod
    def _paragraphs_in_block_range(
        document: DocumentType,
        start_block: int | None,
        end_block: int | None,
    ):
        blocks = list(document.element.body.iterchildren())
        start = start_block if isinstance(start_block, int) else 0
        end = end_block if isinstance(end_block, int) else len(blocks)
        for block in blocks[start:end]:
            if block.tag == qn("w:p"):
                yield Paragraph(block, document._body)
            elif block.tag == qn("w:tbl"):
                table = Table(block, document._body)
                for row in table.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs

    @staticmethod
    def _run_style_snapshot(run, paragraph: Paragraph) -> dict[str, Any]:
        fonts = run._r.rPr.rFonts if run._r.rPr is not None else None
        style: dict[str, Any] = {}
        if fonts is not None:
            for key in (
                "ascii", "hAnsi", "eastAsia", "cs",
                "asciiTheme", "hAnsiTheme", "eastAsiaTheme", "csTheme",
            ):
                value = fonts.get(qn(f"w:{key}"))
                if value:
                    style[key] = value
        try:
            paragraph_style = paragraph.style
        except (AttributeError, KeyError):
            paragraph_style = None
        style_fonts = (
            paragraph_style._element.rPr.rFonts
            if paragraph_style is not None
            and paragraph_style._element.rPr is not None
            else None
        )
        if style_fonts is not None:
            for key in (
                "ascii", "hAnsi", "eastAsia", "cs",
                "asciiTheme", "hAnsiTheme", "eastAsiaTheme", "csTheme",
            ):
                value = style_fonts.get(qn(f"w:{key}"))
                if value:
                    style.setdefault(key, value)
        inherited_name = run.font.name or (
            paragraph_style.font.name if paragraph_style is not None else None
        )
        if inherited_name:
            for key in ("ascii", "hAnsi", "eastAsia", "cs"):
                style.setdefault(key, inherited_name)
        inherited_size = run.font.size or (
            paragraph_style.font.size if paragraph_style is not None else None
        )
        if inherited_size is not None:
            style["size_half_points"] = int(round(inherited_size.pt * 2))
        inherited_bold = run.bold
        if inherited_bold is None and paragraph_style is not None:
            inherited_bold = paragraph_style.font.bold
        if inherited_bold is not None:
            style["bold"] = bool(inherited_bold)
        inherited_italic = run.italic
        if inherited_italic is None and paragraph_style is not None:
            inherited_italic = paragraph_style.font.italic
        if inherited_italic is not None:
            style["italic"] = bool(inherited_italic)
        return style

    @classmethod
    def _extract_fillable_fields(
        cls,
        document: DocumentType,
        start_block: int | None,
        end_block: int | None,
    ) -> tuple[tuple[dict[str, Any], ...], tuple[dict[str, Any], ...]]:
        if start_block is None:
            return (), ()
        fields: list[dict[str, Any]] = []
        actions: list[dict[str, Any]] = []
        seen: set[tuple[str, str]] = set()
        field_key_counts: Counter[str] = Counter()

        def append(
            label: str,
            location: str,
            *,
            surrounding_text: str,
            document_section: str | None,
            table_index: int | None = None,
            paragraph_index: int | None = None,
            row: int | None = None,
            column: int | None = None,
            target_mode: str = "empty_cell",
        ) -> None:
            display_label = re.sub(r"\s+", " ", label).strip().strip("：:")
            semantic_label = re.split(
                r"(?:[_＿]{2,}|…+|\.{3,})", display_label, maxsplit=1
            )[0]
            clean = re.sub(r"[\s()（）]", "", semantic_label)
            if not 2 <= len(clean) <= 40:
                return
            if re.search(r"^甲方$|^乙方$", clean):
                return
            if re.search(r"身份证明|复印件|扫描件|护照|附件", clean):
                return
            if not cls._looks_like_field_label(clean):
                return
            canonical = None
            candidates = [
                (len(alias), -clean.find(alias), key)
                for key, aliases in FIELD_ALIASES.items()
                for alias in aliases
                if alias in clean
            ]
            if candidates:
                canonical = max(candidates)[2]
            slot = SlotContextClassifier.classify(
                label=display_label,
                surrounding_text=surrounding_text,
                source_location=location,
                document_section=document_section,
                table_index=table_index,
                paragraph_index=paragraph_index,
                row=row,
                column=column,
                canonical_hint=canonical,
            )
            if slot.fill_strategy.value == "action_only":
                action_labels = slot.required_actions or (slot.display_name,)
                for action_index, action_label in enumerate(action_labels, start=1):
                    actions.append({
                        "action_id": f"{slot.slot_id}:{action_index}",
                        "display_name": action_label,
                        "source_location": location,
                        "surrounding_text": slot.surrounding_text,
                        "relation_path": [
                            "当前项目", "投标文件", action_label,
                        ],
                        "required_actions": [action_label],
                    })
                return
            canonical_key = slot.canonical_key
            identity = (canonical_key, location)
            if identity in seen:
                return
            seen.add(identity)
            field_key_counts[canonical_key] += 1
            field_key = (
                canonical_key
                if field_key_counts[canonical_key] == 1
                else f"{canonical_key}__{slot.slot_id}"
            )
            source_type = "manual_input"
            if canonical_key in {
                "project_name", "project_number", "project_reference",
                "bid_response_signing_date",
            }:
                source_type = "tender_document"
            elif re.search(r"报价|金额|价格", clean):
                source_type = "pricing_database"
            elif re.search(r"代表|联系人|证书|身份证", clean):
                source_type = "controlled_personnel_vault"
            elif re.search(r"供应商|投标人|地址|电话|邮箱|账号|开户行", clean):
                source_type = "company_profile"
            fields.append({
                "field_key": field_key,
                "canonical_key": canonical_key,
                "label": display_label,
                "required": not bool(
                    re.search(r"可选|如有|若有", clean)
                    or clean in {"备注", "偏离说明"}
                ),
                "expected_source": source_type,
                "source_location": location,
                "target_mode": target_mode,
                **slot.snapshot(),
            })

        blocks = list(document.element.body.iterchildren())
        table_index = 0
        paragraph_index = 0
        current_section: str | None = None
        section_stack: dict[int, str] = {}
        for block in blocks[start_block:end_block]:
            if block.tag.endswith("}tbl"):
                table_index += 1
                table = Table(block, document)
                for row_index, row in enumerate(table.rows, start=1):
                    cell_texts = [cell.text.strip() for cell in row.cells]
                    seen_cell_nodes: set[int] = set()
                    for cell_index, cell in enumerate(row.cells, start=1):
                        cell_node_id = id(cell._tc)
                        if cell_node_id in seen_cell_nodes:
                            continue
                        seen_cell_nodes.add(cell_node_id)
                        target = cell_texts[cell_index - 1]
                        location = (
                            f"表格{table_index}/第{row_index}行/第{cell_index}列"
                        )
                        # Empty table cells are real slots when a nearby label
                        # tells us what belongs there. Support both horizontal
                        # label/value pairs and vertical header/value layouts.
                        if cls._is_blank_target(target):
                            label = cls._table_slot_label(
                                table, row_index, cell_index
                            )
                            if label:
                                append(
                                    label,
                                    location,
                                    surrounding_text=cls._table_slot_context(
                                        table, row_index, cell_index, label
                                    ),
                                    document_section=current_section,
                                    table_index=table_index,
                                    row=row_index,
                                    column=cell_index,
                                    target_mode="empty_cell",
                                )
                                continue

                        # A form cell may contain its own label and blank, e.g.
                        # “法定代表人姓名：____”. Treat the blank as the target,
                        # not the whole cell or the underscore characters.
                        for paragraph_in_cell in cell.paragraphs:
                            cell_line = paragraph_in_cell.text.strip()
                            matches = list(LABELED_BLANK_RE.finditer(cell_line))
                            matches.extend(PAREN_LABEL_BLANK_RE.finditer(cell_line))
                            matches.sort(key=lambda item: item.start())
                            for slot_index, labeled in enumerate(matches, start=1):
                                label = cls._clean_labeled_field(
                                    labeled.group("label")
                                )
                                if not label:
                                    continue
                                append(
                                    label,
                                    f"{location}/第{slot_index}个字段",
                                    surrounding_text=cls._marked_context(
                                        cell_line, labeled.start(), labeled.end(), label
                                    ),
                                    document_section=current_section,
                                    table_index=table_index,
                                    row=row_index,
                                    column=cell_index,
                                    target_mode="inline_cell",
                                )
            elif block.tag.endswith("}p"):
                paragraph_index += 1
                paragraph = Paragraph(block, document)
                text = paragraph.text.strip()
                heading_level, heading_source = cls._heading_level(
                    paragraph, text
                )
                compact_heading = re.sub(r"\s+", "", text)
                if (
                    heading_source == "numbering"
                    and re.match(r"^[一二三四五六七八九十百]+、", compact_heading)
                    and 2 in section_stack
                ):
                    heading_level = 3
                if heading_source == "form_numbering":
                    section_stack = {
                        level: title for level, title in section_stack.items()
                        if level == 1
                    }
                is_structural_heading = cls._is_structural_heading(
                    text, heading_level, heading_source
                )
                if is_structural_heading:
                    assert heading_level is not None
                    section_stack = {
                        level: title for level, title in section_stack.items()
                        if level < heading_level
                    }
                    section_stack[heading_level] = cls._display_heading(text)
                    current_section = " / ".join(
                        section_stack[level] for level in sorted(section_stack)
                    )
                if is_structural_heading:
                    continue
                labeled_matches = list(LABELED_BLANK_RE.finditer(text))
                labeled_matches.extend(PAREN_LABEL_BLANK_RE.finditer(text))
                labeled_matches.sort(key=lambda item: item.start())
                for slot_index, labeled in enumerate(
                    labeled_matches, start=1
                ):
                    label = cls._clean_labeled_field(
                        labeled.group("label")
                    )
                    if not label:
                        continue
                    append(
                        label,
                        f"第{paragraph_index}段/第{slot_index}个字段",
                        surrounding_text=cls._marked_context(
                            text, labeled.start(), labeled.end(), label
                        ),
                        document_section=current_section,
                        paragraph_index=paragraph_index,
                        target_mode="inline_paragraph",
                    )
                inline_matches = list(re.finditer(
                    r"(?:[_＿]{2,}|…+|\.{3,})\s*[（(]([^）)]{1,30})[）)]",
                    text,
                ))
                for slot_index, inline in enumerate(inline_matches, start=1):
                    local_context = cls._local_slot_context(
                        text, inline_matches, slot_index - 1
                    )
                    append(
                        inline.group(1),
                        f"第{paragraph_index}段/第{slot_index}个空位",
                        surrounding_text=local_context,
                        document_section=current_section,
                        paragraph_index=paragraph_index,
                        target_mode="inline_paragraph",
                    )
        unique_actions = {
            item["action_id"]: item for item in actions
        }
        return tuple(fields), tuple(unique_actions.values())

    @staticmethod
    def _local_slot_context(
        text: str,
        matches: list[re.Match[str]],
        index: int,
    ) -> str:
        """Cut one placeholder's clause away from neighbouring placeholders."""
        current = matches[index]
        left = 0 if index == 0 else (
            matches[index - 1].end() + current.start()
        ) // 2
        right = len(text) if index == len(matches) - 1 else (
            current.end() + matches[index + 1].start()
        ) // 2
        return ResponseTemplateService._marked_context(
            text, current.start(), current.end(), current.group(1),
            left=max(0, left - 28),
            right=min(len(text), right + 48),
        )

    @staticmethod
    def _marked_context(
        text: str,
        start: int,
        end: int,
        label: str,
        *,
        left: int | None = None,
        right: int | None = None,
    ) -> str:
        """Preserve local syntax while marking exactly which blank is active."""
        left = max(0, start - 40) if left is None else left
        right = min(len(text), end + 64) if right is None else right
        return (
            text[left:start]
            + f"【当前空位：{label}】"
            + text[end:right]
        )

    @staticmethod
    def _clean_labeled_field(label: str) -> str:
        cleaned = re.sub(r"\s+", " ", label).strip()
        cleaned = re.split(r"[。；;，,]", cleaned)[-1].strip()
        return cleaned

    @staticmethod
    def _is_blank_target(value: str) -> bool:
        text = value.strip()
        return (
            not text
            or bool(PLACEHOLDER_RE.search(text))
            or bool(BLANK_TOKEN_RE.search(text))
        )

    @staticmethod
    def _looks_like_field_label(value: str) -> bool:
        """Reject layout noise while retaining unknown but explicit form fields."""
        compact = re.sub(r"[：:_＿()（）\s]", "", value)
        if not 2 <= len(compact) <= 40:
            return False
        if re.search(r"身份证明|复印件|扫描件|护照|附件", compact):
            return False
        if re.search(r"请|应当|须|不得|^(?:说明|注)$|注[：:]", compact):
            return False
        if re.search(
            r"如下$|承诺$|具备.*条件$|详见|根据.*要求$|参与.*项目$",
            compact,
        ):
            return False
        known = re.search(
            r"名称|姓名|编号|代表|日期|时间|金额|报价|地址|电话|手机|邮箱|"
            r"联系人|职务|职位|岗位|身份证号|证书|资质|签字|盖章|账号|开户行|"
            r"工期|期限|性别|年龄|传真|邮编|网址|法定|代理人|负责人",
            compact,
        )
        # Explicit short labels are retained even before the ontology knows
        # their domain type; they are surfaced as unresolved for review rather
        # than silently dropped.
        return bool(known) or (
            len(compact) <= 12
            and not re.search(r"[。；，,;！？!?]", compact)
        )

    @classmethod
    def _table_slot_label(
        cls,
        table: Table,
        row_index: int,
        column_index: int,
    ) -> str | None:
        row = table.rows[row_index - 1]
        # A horizontal label/value pair must be adjacent. Looking farther left
        # makes every layout cell in a wide table look like the same field.
        if column_index > 1:
            candidate = row.cells[column_index - 2].text.strip()
            if (
                candidate
                and not BLANK_TOKEN_RE.search(candidate)
                and cls._looks_like_field_label(candidate)
            ):
                generic = SlotContextClassifier._generic_business_slot(
                    candidate
                )
                # In a personnel grid the first column is a row category
                # (for example "管理人员"), not the value expected by the
                # next cell.  Continue upward to the real column header.
                if not generic or generic[1] != "project.team.members":
                    return candidate
        # A real grid header applies to all following blank data rows, not just
        # the first one. Propagate it only when its row contains at least two
        # semantic column labels; this avoids carrying arbitrary prose through
        # layout tables.
        for above in range(row_index - 2, -1, -1):
            candidate = table.rows[above].cells[column_index - 1].text.strip()
            if not candidate:
                continue
            if (
                not BLANK_TOKEN_RE.search(candidate)
                and cls._looks_like_field_label(candidate)
                and cls._semantic_label_count(table.rows[above]) >= 2
            ):
                return candidate
            break
        return None

    @classmethod
    def _semantic_label_count(cls, row) -> int:
        labels: set[str] = set()
        seen_nodes: set[int] = set()
        for cell in row.cells:
            node_id = id(cell._tc)
            if node_id in seen_nodes:
                continue
            seen_nodes.add(node_id)
            text = cell.text.strip()
            if (
                text
                and not BLANK_TOKEN_RE.search(text)
                and cls._looks_like_field_label(text)
            ):
                labels.add(re.sub(r"\s+", "", text))
        return len(labels)

    @classmethod
    def _table_slot_context(
        cls,
        table: Table,
        row_index: int,
        column_index: int,
        label: str,
    ) -> str:
        row_text = " | ".join(
            cell.text.strip() for cell in table.rows[row_index - 1].cells
            if cell.text.strip()
        )
        header_text = ""
        if row_index > 1:
            header_text = table.rows[row_index - 2].cells[column_index - 1].text.strip()
        context = " | ".join(item for item in (header_text, row_text) if item)
        return f"{context} | 【当前空位：{label}】" if context else f"【当前空位：{label}】"

    @staticmethod
    def _is_structural_heading(
        title: str,
        level: int | None,
        source: str,
    ) -> bool:
        if level is None:
            return False
        compact = re.sub(r"\s+", "", title)
        if source == "paragraph_style":
            return len(compact) <= 120
        return (
            len(compact) <= 60
            and not BLANK_TOKEN_RE.search(compact)
            and not re.search(r"[。；，]$", compact)
        )

    @classmethod
    def _extract_outline(
        cls,
        document: DocumentType,
        start_block: int | None,
    ) -> tuple[dict[str, Any], ...]:
        """Extract an ordered, user-facing outline up to five levels deep."""
        if start_block is None:
            return ()
        items: list[dict[str, Any]] = []
        seen_form_heading = False
        blocks = list(document.element.body.iterchildren())
        for block_index, block in enumerate(blocks[start_block:], start=start_block):
            if not block.tag.endswith("}p"):
                continue
            paragraph = Paragraph(block, document)
            title = re.sub(r"\s+", " ", paragraph.text).strip()
            if not title or len(title) > 160:
                continue
            compact = title.replace(" ", "")
            if (
                block_index > start_block
                and re.match(
                    r"^第[一二三四五六七八九十百零〇\d]+章",
                    compact,
                )
            ):
                break
            if title.endswith(("。", "；", ";")):
                continue
            level, source = cls._heading_level(paragraph, title)
            if level is None:
                continue
            if source == "form_numbering":
                seen_form_heading = True
            elif (
                seen_form_heading
                and source == "numbering"
                and re.match(
                    r"^[一二三四五六七八九十百]+、", compact
                )
            ):
                level = 3
            items.append(
                {
                    "title": cls._display_heading(title),
                    "source_title": title,
                    "level": min(5, max(1, level)),
                    "order": len(items) + 1,
                    "source": source,
                    "source_block": block_index,
                }
            )
        return tuple(items)

    @staticmethod
    def _display_heading(title: str) -> str:
        """Remove Word form-control numbering noise from user-facing titles."""
        cleaned = re.sub(
            r"^\s*\d+[.．、]?\s*格式\s*\d+\s*",
            "",
            title,
        ).strip()
        return cleaned or title

    @staticmethod
    def _heading_level(
        paragraph: Paragraph,
        title: str,
    ) -> tuple[int | None, str]:
        style_name = ""
        try:
            style_name = paragraph.style.name or ""
        except (AttributeError, KeyError):
            style_name = ""
        style_match = re.search(r"(?:Heading|标题)\s*([1-5])", style_name, re.I)
        if style_match:
            return int(style_match.group(1)), "paragraph_style"
        compact = title.replace(" ", "")
        if re.match(r"^第[一二三四五六七八九十百零〇\d]+章", compact):
            return 1, "numbering"
        if re.match(r"^第[一二三四五六七八九十百零〇\d]+节", compact):
            return 2, "numbering"
        if re.match(r"^格式\s*\d+", compact, re.I):
            return 2, "form_numbering"
        decimal = re.match(r"^(\d+(?:[.．]\d+){1,4})(?:[\s、.]|$)", compact)
        if decimal:
            return decimal.group(1).replace("．", ".").count(".") + 1, "numbering"
        if re.match(r"^[一二三四五六七八九十百]+、", compact):
            return 1, "numbering"
        if re.match(r"^[（(][一二三四五六七八九十百]+[）)]", compact):
            return 2, "numbering"
        if re.match(r"^\d+[、.]", compact):
            return 2, "numbering"
        if re.match(r"^[（(]\d+[）)]", compact):
            return 3, "numbering"
        if re.match(r"^[①②③④⑤⑥⑦⑧⑨⑩]", compact):
            return 4, "numbering"
        return None, "none"

    def fill_docx(
        self,
        *,
        template_content: bytes,
        output_path: Path,
        descriptor: TemplateDescriptor | dict[str, Any],
        field_values: dict[str, str],
        sections: list[dict[str, str]],
        document_title: str | None = None,
    ) -> TemplateFillReport:
        data = (
            descriptor.snapshot()
            if isinstance(descriptor, TemplateDescriptor)
            else descriptor
        )
        if not data.get("detected") or data.get("source_format") != "docx":
            raise ValueError("当前文件不是可自动回填的 DOCX 响应模板。")
        document = Document(BytesIO(template_content))
        if document_title:
            document.core_properties.title = document_title
        start = data.get("start_block")
        end = data.get("end_block")
        if isinstance(start, int) and (
            start > 0 or isinstance(end, int)
        ):
            self._retain_block_range(
                document,
                start,
                end if isinstance(end, int) else None,
            )

        normalized_values = self._normalized_values(field_values)
        descriptor_fields = tuple(data.get("fields") or ())
        descriptor_keys = {
            str(item.get("canonical_key") or item.get("field_key"))
            for item in descriptor_fields
            if item.get("field_key")
        }
        canonical_counts = Counter(
            str(item.get("canonical_key") or item.get("field_key"))
            for item in descriptor_fields
            if item.get("field_key")
        )
        placeholder_values = {
            key: value for key, value in normalized_values.items()
            if canonical_counts.get(key, 0) <= 1
        }
        filled: set[str] = set()
        unresolved: set[str] = set()
        for paragraph in self._all_paragraphs(document):
            replaced, used, missing = self._replace_placeholders(
                paragraph.text, placeholder_values
            )
            if replaced != paragraph.text:
                self._replace_paragraph_text(paragraph, replaced)
            filled.update(used)
            unresolved.update(missing)
        self._fill_descriptor_fields(
            document,
            descriptor_fields,
            normalized_values,
            filled,
        )
        legacy_values = {
            key: value for key, value in normalized_values.items()
            if key not in descriptor_keys
        }
        self._fill_label_paragraphs(document, legacy_values, filled)
        self._fill_label_cells(document, legacy_values, filled)
        self._replace_legacy_x_placeholders(document, legacy_values, filled)
        descriptor_labels = set(data.get("field_labels") or ())
        for key, aliases in FIELD_ALIASES.items():
            if any(alias in descriptor_labels for alias in aliases):
                if key not in normalized_values:
                    unresolved.add(key)

        inserted: list[str] = []
        missing_sections: list[str] = []
        for section in sections:
            title = section["title"].strip()
            content = section.get("content", "").strip()
            target = self._find_heading(document, title)
            if target is None:
                missing_sections.append(title)
                continue
            self._insert_content_after(
                target,
                content,
                font_profile=data.get("font_profile") or {},
            )
            inserted.append(title)

        if missing_sections:
            generic_target = self._find_heading(document, "技术方案")
            if generic_target is not None:
                combined = "\n".join(
                    f"## {item['title']}\n{item.get('content', '')}"
                    for item in sections
                    if item["title"] in missing_sections
                )
                self._insert_content_after(
                    generic_target,
                    combined,
                    font_profile=data.get("font_profile") or {},
                )
                inserted.extend(missing_sections)
                missing_sections = []

        # Keep the template's directory exactly where the source placed it.
        # Word refreshes any existing TOC/PAGEREF fields on open; we never
        # create a second generic directory in strict-template mode.
        self._request_field_refresh(document)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
        return TemplateFillReport(
            filled_fields=tuple(sorted(filled)),
            unresolved_fields=tuple(sorted(unresolved - filled)),
            inserted_sections=tuple(inserted),
            unresolved_sections=tuple(missing_sections),
        )

    @staticmethod
    def _request_field_refresh(document: DocumentType) -> None:
        settings = document.settings.element
        update_fields = settings.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            settings.append(update_fields)
        update_fields.set(qn("w:val"), "true")

    @staticmethod
    def _block_text(block) -> str:
        # python-docx exposes aggregate text on container nodes as well as on
        # their child ``w:t`` nodes. Walking every element therefore repeats
        # the same heading two or three times and makes boundary ranking depend
        # on XML shape. Read only Word text leaves.
        return "".join(
            node.text or ""
            for node in block.xpath(".//*[local-name()='t']")
        ).strip()

    @staticmethod
    def _retain_block_range(
        document: DocumentType,
        start: int,
        end: int | None,
    ) -> None:
        body = document.element.body
        blocks = list(body.iterchildren())
        for index, block in enumerate(blocks):
            if block.tag.endswith("}sectPr"):
                continue
            if index < start or (end is not None and index >= end):
                body.remove(block)

    @staticmethod
    def _all_paragraphs(document: DocumentType):
        yield from document.paragraphs
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs

    @staticmethod
    def _normalized_values(values: dict[str, str]) -> dict[str, str]:
        result = {
            key: str(value).strip()
            for key, value in values.items()
            if value is not None and str(value).strip()
        }
        for key, aliases in FIELD_ALIASES.items():
            if key not in result:
                continue
            for alias in aliases:
                result.setdefault(alias, result[key])
        return result

    @staticmethod
    def _replace_placeholders(
        text: str,
        values: dict[str, str],
    ) -> tuple[str, set[str], set[str]]:
        used: set[str] = set()
        missing: set[str] = set()

        def replace(match: re.Match) -> str:
            key = next(group for group in match.groups() if group).strip()
            value = values.get(key)
            if value is None:
                missing.add(key)
                return match.group(0)
            used.add(key)
            return value

        return PLACEHOLDER_RE.sub(replace, text), used, missing

    @staticmethod
    def _replace_paragraph_text(paragraph: Paragraph, value: str) -> None:
        if paragraph.runs:
            paragraph.runs[0].text = value
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(value)

    @staticmethod
    def _fill_label_cells(
        document: DocumentType,
        values: dict[str, str],
        filled: set[str],
    ) -> None:
        for table in document.tables:
            for row in table.rows:
                for index, cell in enumerate(row.cells[:-1]):
                    label = re.sub(r"[\s:：()（）]", "", cell.text)
                    for key, aliases in FIELD_ALIASES.items():
                        if not any(re.sub(r"[\s:：()（）]", "", alias) in label for alias in aliases):
                            continue
                        value = values.get(key)
                        target = row.cells[index + 1]
                        if value and ResponseTemplateService._is_blank_target(
                            target.text
                        ):
                            ResponseTemplateService._set_cell_text_preserving_style(
                                target, value, source_cell=cell
                            )
                            filled.add(key)
                        break

    @staticmethod
    def _fill_descriptor_fields(
        document: DocumentType,
        fields: list[dict[str, Any]] | tuple[dict[str, Any], ...],
        values: dict[str, str],
        filled: set[str],
    ) -> None:
        unresolved_entries: list[tuple[str, str]] = []
        for item in fields:
            key = str(item.get("field_key") or "")
            value = values.get(key)
            if not key or not value:
                continue
            table_index = item.get("table_index")
            row_index = item.get("row")
            column_index = item.get("column")
            if all(
                isinstance(index, int)
                for index in (table_index, row_index, column_index)
            ):
                try:
                    table = document.tables[table_index - 1]
                    row = table.rows[row_index - 1]
                    target = row.cells[column_index - 1]
                    source = row.cells[max(0, column_index - 2)]
                except (IndexError, TypeError):
                    unresolved_entries.append((str(item.get("label") or ""), key))
                    continue
                if item.get("target_mode") == "inline_cell":
                    replaced = False
                    for paragraph in target.paragraphs:
                        updated, count = ResponseTemplateService._replace_text_slot(
                            paragraph.text,
                            str(item.get("label") or ""),
                            value,
                        )
                        if count:
                            ResponseTemplateService._replace_paragraph_text(
                                paragraph, updated
                            )
                            filled.add(key)
                            replaced = True
                            break
                    if not replaced:
                        unresolved_entries.append(
                            (str(item.get("label") or ""), key)
                        )
                elif ResponseTemplateService._is_blank_target(target.text):
                    ResponseTemplateService._set_cell_text_preserving_style(
                        target, value, source_cell=source
                    )
                    filled.add(key)
                continue
            paragraph_index = item.get("paragraph_index")
            if isinstance(paragraph_index, int):
                try:
                    paragraph = document.paragraphs[paragraph_index - 1]
                except IndexError:
                    unresolved_entries.append((str(item.get("label") or ""), key))
                    continue
                original = paragraph.text
                label = str(item.get("label") or "").strip()
                updated, count = ResponseTemplateService._replace_text_slot(
                    original, label, value
                )
                if count:
                    ResponseTemplateService._replace_paragraph_text(
                        paragraph, updated
                    )
                    filled.add(key)
                continue
            unresolved_entries.append((str(item.get("label") or ""), key))

        # Backward compatibility for descriptors created before slot coordinates.
        for table in document.tables:
            for row in table.rows:
                for index, cell in enumerate(row.cells[:-1]):
                    label_text = re.sub(r"[\s:：()（）]", "", cell.text)
                    match = next(
                        (
                            (label, key) for label, key in unresolved_entries
                            if re.sub(r"[\s:：()（）]", "", label) == label_text
                            and values.get(key)
                        ),
                        None,
                    )
                    if match is None:
                        continue
                    _, key = match
                    target = row.cells[index + 1]
                    if not target.text.strip() or PLACEHOLDER_RE.search(target.text):
                        ResponseTemplateService._set_cell_text_preserving_style(
                            target, values[key], source_cell=cell
                        )
                        filled.add(key)
        for paragraph in ResponseTemplateService._all_paragraphs(document):
            original = paragraph.text
            updated = original
            for label, key in unresolved_entries:
                value = values.get(key)
                if not value:
                    continue
                pattern = re.compile(
                    rf"^(?P<label>{re.escape(label)}\s*[:：])"
                    r"[ \t]*(?:[_＿]{2,}|…+|\.{3,})?[ \t]*$"
                )
                updated, count = pattern.subn(
                    lambda match: f"{match.group('label')} {value}", updated
                )
                if count:
                    filled.add(key)
            if updated != original:
                ResponseTemplateService._replace_paragraph_text(paragraph, updated)

    @staticmethod
    def _replace_text_slot(
        text: str,
        label: str,
        value: str,
    ) -> tuple[str, int]:
        """Fill exactly one semantic blank while preserving its label text."""
        escaped = re.escape(label.strip())
        patterns = (
            # 姓名：____ / 姓名： (end of paragraph or cell)
            re.compile(
                rf"(?P<prefix>{escaped}\s*[:：]\s*)"
                r"(?:[_＿]{1,}|…{2,}|\.{3,})?(?=$|\s)"
            ),
            # （姓名）____
            re.compile(
                rf"(?P<prefix>[（(]{escaped}[）)]\s*)"
                r"(?:[_＿]{1,}|…{2,}|\.{3,})"
            ),
            # ____（姓名）
            re.compile(
                rf"(?:[_＿]{{1,}}|…{{2,}}|\.{{3,}})\s*"
                rf"(?P<suffix>[（(]{escaped}[）)])"
            ),
        )
        for index, pattern in enumerate(patterns):
            if index == 2:
                updated, count = pattern.subn(
                    lambda match: f"{value}{match.group('suffix')}",
                    text,
                    count=1,
                )
            else:
                updated, count = pattern.subn(
                    lambda match: f"{match.group('prefix')}{value}",
                    text,
                    count=1,
                )
            if count:
                return updated, count
        return text, 0

    @staticmethod
    def _fill_label_paragraphs(
        document: DocumentType,
        values: dict[str, str],
        filled: set[str],
    ) -> None:
        label_patterns = {
            "project_name": r"(?:项目名称|采购项目名称|招标项目名称)",
            "project_number": r"(?:项目编号|采购编号|招标编号)",
            "bidder_name": r"(?:供应商名称|投标人名称|响应人名称|供应商)",
        }
        for paragraph in ResponseTemplateService._all_paragraphs(document):
            original = paragraph.text
            updated = original
            for key, label_pattern in label_patterns.items():
                value = values.get(key)
                if not value:
                    continue
                pattern = re.compile(
                    rf"(?P<label>{label_pattern}(?:\s*[（(][^）)]{{0,24}}[）)])?\s*[:：])"
                    r"(?P<blank>[ \t]*(?:[_＿]{2,}|…+|\.{3,})?[ \t]*$)"
                )
                updated, count = pattern.subn(
                    lambda match: f"{match.group('label')} {value}", updated
                )
                if count:
                    filled.add(key)
            if updated != original:
                ResponseTemplateService._replace_paragraph_text(paragraph, updated)

    @staticmethod
    def _replace_legacy_x_placeholders(
        document: DocumentType,
        values: dict[str, str],
        filled: set[str],
    ) -> None:
        """Replace legacy XXX blanks without flattening the source layout."""
        for paragraph in ResponseTemplateService._all_paragraphs(document):
            original = paragraph.text
            updated = original
            for key, aliases in FIELD_ALIASES.items():
                value = values.get(key)
                if not value:
                    continue
                alias_pattern = "|".join(
                    sorted((re.escape(alias) for alias in aliases), key=len, reverse=True)
                )
                pattern = re.compile(
                    rf"(?P<label>(?:{alias_pattern})(?:\s*[（(][^）)]{{0,24}}[）)])?\s*[:：]?)"
                    r"\s*[Xx]{2,}"
                )
                updated, count = pattern.subn(
                    lambda match: f"{match.group('label')} {value}", updated
                )
                if count:
                    filled.add(key)
            updated = re.sub(
                r"[Xx]{2,4}年[Xx]{1,4}月[Xx]{1,4}日",
                "【待审核日期】",
                updated,
            )
            updated = re.sub(r"[Xx]{3,}", "【待审核】", updated)
            if updated != original:
                ResponseTemplateService._replace_paragraph_text(paragraph, updated)

    @staticmethod
    def _fillable_labels(document: DocumentType) -> tuple[str, ...]:
        labels: set[str] = set()
        for table in document.tables:
            for row in table.rows:
                for index, cell in enumerate(row.cells[:-1]):
                    label = re.sub(r"[\s:：()（）]", "", cell.text)
                    target = row.cells[index + 1].text.strip()
                    if target and not PLACEHOLDER_RE.search(target):
                        continue
                    for aliases in FIELD_ALIASES.values():
                        match = next(
                            (
                                alias for alias in aliases
                                if re.sub(r"[\s:：()（）]", "", alias)
                                in label
                            ),
                            None,
                        )
                        if match:
                            labels.add(match)
                            break
        return tuple(sorted(labels))

    @staticmethod
    def _find_heading(document: DocumentType, title: str) -> Paragraph | None:
        normalized = re.sub(r"[\s\d一二三四五六七八九十、.．()（）]", "", title)
        aliases = {
            "技术方案": ("技术方案", "项目实施方案", "服务方案"),
            "项目实施方案": ("项目实施方案", "技术方案", "服务方案"),
        }
        targets = aliases.get(normalized, (normalized,))
        paragraphs = list(ResponseTemplateService._all_paragraphs(document))
        for paragraph in paragraphs:
            candidate = re.sub(
                r"[\s\d一二三四五六七八九十、.．()（）]", "", paragraph.text
            )
            if candidate in targets:
                return paragraph
        for paragraph in paragraphs:
            candidate = re.sub(
                r"[\s\d一二三四五六七八九十、.．()（）]", "", paragraph.text
            )
            if any(len(target) >= 4 and target in candidate for target in targets):
                return paragraph
        return None

    @staticmethod
    def _insert_content_after(
        paragraph: Paragraph,
        content: str,
        *,
        font_profile: dict[str, Any] | None = None,
    ) -> None:
        anchor = paragraph._p
        profile = font_profile or {}
        for raw_line in reversed([line.strip() for line in content.splitlines() if line.strip()]):
            new_p = OxmlElement("w:p")
            heading = bool(re.match(r"^\d+(?:[.．]\d+)*[.、．]?\s+\S", raw_line))
            properties = OxmlElement("w:pPr")
            spacing = OxmlElement("w:spacing")
            spacing.set(qn("w:line"), "360")
            spacing.set(qn("w:lineRule"), "auto")
            spacing.set(qn("w:before"), "160" if heading else "0")
            spacing.set(qn("w:after"), "100")
            properties.append(spacing)
            if not heading:
                indentation = OxmlElement("w:ind")
                indentation.set(qn("w:firstLineChars"), "200")
                properties.append(indentation)
            new_p.append(properties)
            run = OxmlElement("w:r")
            run_properties = OxmlElement("w:rPr")
            style = profile.get("heading" if heading else "body") or {}
            ResponseTemplateService._apply_run_style(run_properties, style)
            color = OxmlElement("w:color")
            color.set(qn("w:val"), "000000")
            run_properties.append(color)
            if heading:
                run_properties.append(OxmlElement("w:b"))
            run.append(run_properties)
            text = OxmlElement("w:t")
            text.text = re.sub(r"^#{1,6}\s*", "", raw_line)
            run.append(text)
            new_p.append(run)
            anchor.addnext(new_p)

    @staticmethod
    def _apply_run_style(run_properties, style: dict[str, Any]) -> None:
        font_values = {
            key: style.get(key)
            for key in (
                "ascii", "hAnsi", "eastAsia", "cs",
                "asciiTheme", "hAnsiTheme", "eastAsiaTheme", "csTheme",
            )
            if style.get(key)
        }
        if font_values:
            fonts = OxmlElement("w:rFonts")
            for key, value in font_values.items():
                fonts.set(qn(f"w:{key}"), str(value))
            run_properties.append(fonts)
        size_value = style.get("size_half_points")
        if size_value:
            for tag in ("w:sz", "w:szCs"):
                size = OxmlElement(tag)
                size.set(qn("w:val"), str(size_value))
                run_properties.append(size)
        if style.get("bold"):
            run_properties.append(OxmlElement("w:b"))
        if style.get("italic"):
            run_properties.append(OxmlElement("w:i"))

    @staticmethod
    def _set_cell_text_preserving_style(
        target,
        value: str,
        *,
        source_cell=None,
    ) -> None:
        paragraph = target.paragraphs[0]
        if paragraph.runs:
            paragraph.runs[0].text = value
            for run in paragraph.runs[1:]:
                run.text = ""
            return
        run = paragraph.add_run(value)
        if source_cell is None:
            return
        source_run = next(
            (
                item
                for item in source_cell.paragraphs[0].runs
                if item._r.rPr is not None
            ),
            None,
        )
        if source_run is not None:
            run._r.insert(0, deepcopy(source_run._r.rPr))
