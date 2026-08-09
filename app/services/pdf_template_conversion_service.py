from __future__ import annotations

import subprocess
import sys
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from pypdf import PdfReader


_CONVERSION_SCRIPT = """
from pdf2docx import Converter
import sys
converter = Converter(sys.argv[1])
try:
    converter.convert(sys.argv[2])
finally:
    converter.close()
"""


@dataclass(frozen=True)
class PdfConversionResult:
    status: str
    content: bytes | None
    page_count: int
    paragraph_count: int
    table_count: int
    message: str

    def snapshot(self) -> dict[str, object]:
        return {
            "status": self.status,
            "page_count": self.page_count,
            "paragraph_count": self.paragraph_count,
            "table_count": self.table_count,
            "message": self.message,
        }


class PdfTemplateConversionService:
    """Convert PDF to editable DOCX locally before choosing a writer."""

    def __init__(self, timeout_seconds: int = 240):
        self.timeout_seconds = timeout_seconds

    def convert(self, content: bytes) -> PdfConversionResult:
        try:
            page_count = len(PdfReader(BytesIO(content)).pages)
        except Exception as exc:
            return PdfConversionResult(
                "failed", None, 0, 0, 0,
                f"PDF结构无法读取：{type(exc).__name__}",
            )
        with TemporaryDirectory(prefix="bid-pdf-conversion-") as directory:
            source = Path(directory) / "source.pdf"
            target = Path(directory) / "converted.docx"
            source.write_bytes(content)
            try:
                completed = subprocess.run(
                    [
                        sys.executable,
                        "-c",
                        _CONVERSION_SCRIPT,
                        str(source),
                        str(target),
                    ],
                    stdin=subprocess.DEVNULL,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.STDOUT,
                    timeout=self.timeout_seconds,
                    check=False,
                    text=True,
                )
            except subprocess.TimeoutExpired:
                return PdfConversionResult(
                    "failed", None, page_count, 0, 0,
                    "PDF转Word超过安全时限，已停止转换。",
                )
            except OSError as exc:
                return PdfConversionResult(
                    "failed", None, page_count, 0, 0,
                    f"PDF转Word组件无法启动：{type(exc).__name__}",
                )
            if completed.returncode != 0 or not target.is_file():
                return PdfConversionResult(
                    "failed", None, page_count, 0, 0,
                    "PDF转Word失败，未生成可编辑文档。",
                )
            converted = target.read_bytes()
            try:
                document = Document(target)
                paragraphs = sum(
                    bool(item.text.strip()) for item in document.paragraphs
                )
                tables = len(document.tables)
            except Exception as exc:
                return PdfConversionResult(
                    "structure_validation_failed", None, page_count, 0, 0,
                    f"转换文件无法作为Word打开：{type(exc).__name__}",
                )
            if paragraphs == 0 and tables == 0:
                return PdfConversionResult(
                    "structure_validation_failed", None, page_count, 0, 0,
                    "转换后的Word没有可编辑文字或表格。",
                )
            return PdfConversionResult(
                "succeeded", converted, page_count, paragraphs, tables,
                "PDF已转换为可编辑Word并通过基础结构校验。",
            )
